from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.color.rgb = RGBColor(0, 51, 102)
    hs.font.bold = True
    if level == 1:
        hs.font.size = Pt(18)
        hs.paragraph_format.space_before = Pt(24)
        hs.paragraph_format.space_after = Pt(12)
    elif level == 2:
        hs.font.size = Pt(15)
        hs.paragraph_format.space_before = Pt(18)
        hs.paragraph_format.space_after = Pt(8)
    else:
        hs.font.size = Pt(13)
        hs.paragraph_format.space_before = Pt(12)
        hs.paragraph_format.space_after = Pt(6)

def add_title(text, size=24, bold=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0, 51, 102)
    return p

def add_body(text):
    doc.add_paragraph(text)

def add_bullet(text):
    doc.add_paragraph(text, style='List Bullet')

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            table.rows[ri + 1].cells[ci].text = str(val)
    doc.add_paragraph()
    return table

# ============================================
# COVER PAGE
# ============================================
for _ in range(4):
    doc.add_paragraph()

add_title("PADHAANEWALA EDUTECH SERVICES")
add_title("BENGALURU - 560100", size=16, bold=False)
doc.add_paragraph()
add_title("PROJECT PLAN", size=28)
doc.add_paragraph()
add_title("Padhaanewala Education Technology Platform", size=14, bold=False)
add_title("padhaanewala.in", size=12, bold=False)
add_title("Version: 1.0", size=12, bold=False)
add_title("Date: September 2026", size=12, bold=False)

doc.add_page_break()

# ============================================
# 1. PROJECT OVERVIEW
# ============================================
doc.add_heading('1. PROJECT OVERVIEW', level=1)

doc.add_heading('1.1. Project Summary', level=2)
add_body("Padhaanewala is a comprehensive education technology platform designed to help students discover colleges, search courses, compare institutions, find scholarships, take mock tests, and receive AI-powered guidance. The platform is database-driven, scalable, and includes a full admin panel, CRM system, and lead management infrastructure.")

doc.add_heading('1.2. Project Objectives', level=2)
add_bullet("Build a professional, scalable, database-driven education platform")
add_bullet("Provide accurate college and course information to students")
add_bullet("Generate and manage admission leads for Padhaanewala business")
add_bullet("Implement AI-powered features (College Predictor, Education Assistant)")
add_bullet("Ensure admin can manage all content without coding knowledge")
add_bullet("Achieve strong SEO presence for organic traffic growth")

doc.add_heading('1.3. Project Scope', level=2)
add_body("IN SCOPE:")
add_bullet("Student-facing website (Next.js)")
add_bullet("College, Course, Scholarship, Exam databases")
add_bullet("AI College Predictor with RAG")
add_bullet("AI Education Assistant with RAG")
add_bullet("Mock Test system")
add_bullet("College Comparison tool")
add_bullet("Student accounts and dashboards")
add_bullet("Admission enquiry and lead management system")
add_bullet("Admin panel with CMS capabilities")
add_bullet("Counsellor dashboard with role-based access")
add_bullet("SEO infrastructure and programmatic SEO")
add_bullet("Analytics and reporting")
add_bullet("Blog and content management")

add_body("OUT OF SCOPE (Future phases):")
add_bullet("Mobile application (planned for future)")
add_bullet("Online payments (planned for future)")
add_bullet("Premium memberships (planned for future)")
add_bullet("Multilingual support (planned for future)")

doc.add_heading('1.4. Key Stakeholders', level=2)
stakeholders = [
    ("Project Owner", "Padhaanewala Founder", "Approvals, Business decisions, Final sign-off"),
    ("Project Manager", "To be assigned", "Planning, Coordination, Risk management"),
    ("Backend Engineers", "To be assigned", "FastAPI development, Database, APIs"),
    ("Frontend Engineers", "To be assigned", "Next.js development, UI, Components"),
    ("AI/ML Engineer", "To be assigned", "RAG pipeline, Embeddings, AI features"),
    ("DevOps Engineer", "To be assigned", "Deployment, CI/CD, Monitoring, Infrastructure"),
    ("UI/UX Designer", "To be assigned", "Design systems, Wireframes, Prototypes"),
    ("Content Team", "To be assigned", "College data entry, Course content, Scholarships"),
    ("QA Engineers", "To be assigned", "Testing, Quality assurance, Bug tracking"),
]
add_table(["Role", "Person", "Responsibility"], stakeholders)

doc.add_heading('1.5. Project Success Criteria', level=2)
add_bullet("Homepage works correctly and looks professional")
add_bullet("College search returns accurate, filtered results")
add_bullet("College and Course pages display complete information")
add_bullet("Admin can edit content without modifying code")
add_bullet("Enquiries reach admin panel and CRM immediately")
add_bullet("Authentication works securely (JWT + OTP)")
add_bullet("Mobile version works on all devices")
add_bullet("SEO basics configured (sitemap, schema, clean URLs)")
add_bullet("HTTPS enabled and functioning")
add_bullet("Database backup automated and tested")
add_bullet("Security testing completed without critical issues")
add_bullet("All major APIs tested and working")
add_bullet("Error handling works (no technical errors shown to users)")
add_bullet("Analytics configured (Google Analytics)")
add_bullet("Sitemap available at /sitemap.xml")
add_bullet("Production deployment stable")

doc.add_page_break()

# ============================================
# 2. PROJECT TIMELINE
# ============================================
doc.add_heading('2. PROJECT TIMELINE', level=1)

doc.add_heading('2.1. High-Level Timeline', level=2)
timeline = [
    ("Month 1", "Foundation", "Project setup, Database design, Authentication"),
    ("Month 2", "Core Features", "College system, Course system"),
    ("Month 3", "Content & Search", "Scholarships, Exams, Search enhancement"),
    ("Month 4", "Student Features", "Student accounts, Reviews, Comparison"),
    ("Month 5", "AI Features", "RAG setup, AI Predictor, AI Assistant"),
    ("Month 6", "Admin & CMS", "Admin panel, Content management"),
    ("Month 7", "Mock Tests", "Question bank, Test engine"),
    ("Month 8", "CRM & Leads", "Lead management, Counsellor system"),
    ("Month 9", "Testing & Launch", "Testing, Deployment, Production launch"),
]
add_table(["Month", "Phase", "Key Deliverables"], timeline)

doc.add_heading('2.2. Detailed Weekly Breakdown', level=2)

doc.add_heading('Month 1: Foundation (Weeks 1-4)', level=3)
weeks = [
    ("Week 1", "Project initialization, Next.js + FastAPI setup, Docker configuration", "Working dev environment"),
    ("Week 2", "CI/CD pipeline (GitHub Actions), Authentication system (JWT + OTP)", "Dev + Staging environments"),
    ("Week 3", "All 45+ database tables, Alembic migrations, pgvector extension", "Complete database schema"),
    ("Week 4", "Seed data, Environment variables, Basic API endpoints", "Basic API working"),
]
add_table(["Week", "Activities", "Deliverable"], weeks)

doc.add_heading('Month 2: Core Features (Weeks 5-8)', level=3)
weeks = [
    ("Week 5", "College CRUD (Admin API), College pages (Frontend)", "College management system"),
    ("Week 6", "College search with filters, CSV import functionality", "College import + search"),
    ("Week 7", "College gallery upload, Image optimization pipeline", "Gallery working"),
    ("Week 8", "Course CRUD, Course pages, College-Course relationships, Fee structure", "Course system complete"),
]
add_table(["Week", "Activities", "Deliverable"], weeks)

doc.add_heading('Month 3: Content & Search (Weeks 9-12)', level=3)
weeks = [
    ("Week 9", "Scholarship database + pages", "Scholarship system"),
    ("Week 10", "Exam database + pages, Deadline tracking", "Exam system"),
    ("Week 11", "Natural language search (basic), Advanced filter UI", "Search enhancement"),
    ("Week 12", "Search pagination, Suggestions/autocomplete", "Full search working"),
]
add_table(["Week", "Activities", "Deliverable"], weeks)

doc.add_heading('Month 4: Student Features (Weeks 13-16)', level=3)
weeks = [
    ("Week 13", "Registration/Login (OTP + Email), Student dashboard", "Student authentication"),
    ("Week 14", "Saved colleges, Profile management", "Student profile complete"),
    ("Week 15", "Review submission + moderation", "Review system"),
    ("Week 16", "College comparison tool, FAQ system", "Comparison complete"),
]
add_table(["Week", "Activities", "Deliverable"], weeks)

doc.add_heading('Month 5: AI Features (Weeks 17-20)', level=3)
weeks = [
    ("Week 17", "pgvector configuration, Embedding generation service", "RAG foundation"),
    ("Week 18", "document_embeddings table, Vector search API, Celery embedding task", "RAG pipeline"),
    ("Week 19", "AI College Predictor with RAG", "Predictor working"),
    ("Week 20", "AI Education Assistant with RAG, AI Comparison, Source citation", "All AI features"),
]
add_table(["Week", "Activities", "Deliverable"], weeks)

doc.add_heading('Month 6: Admin & CMS (Weeks 21-24)', level=3)
weeks = [
    ("Week 21", "All 18 admin modules", "Admin panel complete"),
    ("Week 22", "Role-based access control, Media management", "Access control"),
    ("Week 23", "Blog system, FAQ management", "CMS content"),
    ("Week 24", "Banner management, SEO metadata management", "Full CMS working"),
]
add_table(["Week", "Activities", "Deliverable"], weeks)

doc.add_heading('Month 7: Mock Tests (Weeks 25-28)', level=3)
weeks = [
    ("Week 25", "Question CRUD (Admin), Question import (CSV)", "Question bank"),
    ("Week 26", "Question categories/difficulty, Question validation", "Question bank ready"),
    ("Week 27", "Test creation (Admin), Test taking interface (Student)", "Test engine"),
    ("Week 28", "Timer + navigation, Submission + scoring, Results + analytics", "Mock test complete"),
]
add_table(["Week", "Activities", "Deliverable"], weeks)

doc.add_heading('Month 8: CRM & Leads (Weeks 29-32)', level=3)
weeks = [
    ("Week 29", "Enquiry forms (all pages), Lead management (Admin)", "Lead capture"),
    ("Week 30", "Lead source tracking, UTM tracking", "Lead attribution"),
    ("Week 31", "Counsellor dashboard, Lead assignment", "Counsellor system"),
    ("Week 32", "Follow-up system, Notes system, Contact status", "Full CRM complete"),
]
add_table(["Week", "Activities", "Deliverable"], weeks)

doc.add_heading('Month 9: Testing & Launch (Weeks 33-36)', level=3)
weeks = [
    ("Week 33", "Unit tests, Integration tests, E2E tests (Playwright)", "Test report"),
    ("Week 34", "Security testing, Performance testing, Mobile testing", "Quality assurance"),
    ("Week 35", "Production server setup, DNS, SSL, CDN, Monitoring, Backup", "Production ready"),
    ("Week 36", "Launch checklist verification, Production launch", "LAUNCHED"),
]
add_table(["Week", "Activities", "Deliverable"], weeks)

doc.add_page_break()

# ============================================
# 3. RESOURCE PLAN
# ============================================
doc.add_heading('3. RESOURCE PLAN', level=1)

doc.add_heading('3.1. Team Requirements', level=2)
team = [
    ("Backend Developer (Python/FastAPI)", "1-2", "Full-time", "APIs, Database, Auth, AI integration"),
    ("Frontend Developer (Next.js/React)", "1-2", "Full-time", "UI, Pages, Components, Responsive"),
    ("Database Designer", "1", "Part-time", "Schema, Migrations, Optimization"),
    ("AI/ML Engineer (RAG)", "1", "Part-time", "Embeddings, Vector search, AI pipeline"),
    ("DevOps Engineer", "1", "Part-time", "Docker, CI/CD, Deployment, Monitoring"),
    ("UI/UX Designer", "1", "Part-time", "Design system, Wireframes, Prototypes"),
    ("Content/Data Entry Team", "2-3", "Full-time (Month 2+)", "College data, Course data, Scholarships"),
    ("QA Engineer", "1", "Part-time (Month 8+)", "Testing, Quality assurance"),
]
add_table(["Role", "Count", "Commitment", "Responsibilities"], team)

doc.add_heading('3.2. Technology Requirements', level=2)
add_bullet("Frontend: Next.js, React, TypeScript, Tailwind CSS")
add_bullet("Backend: Python, FastAPI")
add_bullet("Database: PostgreSQL with pgvector")
add_bullet("Cache: Redis")
add_bullet("Storage: AWS S3 / Cloudflare R2")
add_bullet("Search: PostgreSQL full-text search (later OpenSearch)")
add_bullet("AI: OpenAI API + RAG pipeline")
add_bullet("Queue: Celery")
add_bullet("Containerization: Docker")
add_bullet("CI/CD: GitHub Actions")
add_bullet("Monitoring: Sentry, Datadog")
add_bullet("CDN: Cloudflare")

doc.add_heading('3.3. Infrastructure Requirements', level=2)
infra = [
    ("Development Server", "2 vCPU, 4GB RAM, 50GB SSD", "Stage 1 (Month 1)"),
    ("Staging Server", "2 vCPU, 4GB RAM, 50GB SSD", "Stage 2 (Month 1)"),
    ("Production Server (Min)", "4 vCPU, 8GB RAM, 100GB SSD", "Stage 3 (Month 9)"),
    ("Production Server (Rec)", "8 vCPU, 16GB RAM, 200GB SSD", "Stage 3 (Month 9)"),
    ("Object Storage", "S3/R2 - min 50GB", "Stage 1 (Month 1)"),
    ("Domain", "padhaanewala.in", "Stage 1 (Month 1)"),
]
add_table(["Resource", "Specification", "Required By"], infra)

doc.add_page_break()

# ============================================
# 4. BUDGET PLAN
# ============================================
doc.add_heading('4. BUDGET PLAN', level=1)

doc.add_heading('4.1. Development Cost Estimate', level=2)
add_body("Note: Development cost depends on team size, location, and engagement model. Estimates below are indicative.")

costs = [
    ("Team (7-9 members)", "Monthly", "$8,000 - $15,000"),
    ("Infrastructure", "Monthly", "$160 - $400"),
    ("Third-party services", "Monthly", "$50 - $150"),
    ("Total Monthly (Development)", "", "$8,200 - $15,550"),
    ("Development Duration", "9 months", ""),
    ("Total Development Cost (Est.)", "", "$74,000 - $140,000"),
]
add_table(["Item", "Basis", "Estimate"], costs)

doc.add_heading('4.2. Monthly Operating Cost (Post-Launch)', level=2)
ops = [
    ("Cloud Server (4 vCPU, 8GB)", "AWS/DigitalOcean", "$40-80"),
    ("Cloudflare Pro", "CDN + WAF", "$20"),
    ("AWS S3 / R2", "Storage (100GB)", "$5"),
    ("SendGrid", "Email (1000/day)", "$20"),
    ("MSG91", "SMS", "$25"),
    ("OpenAI API", "AI usage", "$50-100"),
    ("Domain", "padhaanewala.in", "$1/month"),
    ("Sentry", "Error tracking", "$0-26"),
    ("Monitoring", "Uptime/Perf", "$0-50"),
    ("Team Maintenance", "Support", "$2,000-5,000"),
    ("Total (Min)", "", "$2,160"),
    ("Total (Recommended)", "", "$5,350"),
]
add_table(["Item", "Purpose", "Monthly Cost"], ops)

doc.add_heading('4.3. Cost Control Measures', level=2)
add_bullet("Phase-wise development - launch early with core features")
add_bullet("Open source tools wherever possible")
add_bullet("Optimize AI usage - cache common queries")
add_bullet("Start with minimum infrastructure, scale as needed")
add_bullet("Use free tiers for monitoring tools initially")

doc.add_page_break()

# ============================================
# 5. RISK MANAGEMENT PLAN
# ============================================
doc.add_heading('5. RISK MANAGEMENT PLAN', level=1)

doc.add_heading('5.1. Risk Identification and Mitigation', level=2)
risks = [
    ("College Data Accuracy", "High", "Wrong fees/admission data harm trust", "Data verification workflow, official sources, periodic re-verification, user feedback"),
    ("AI Hallucination", "High", "AI invents college names/facts", "RAG pipeline, source citations, strict guardrails, no-hallucination rules"),
    ("Team Capacity", "Medium", "Solo developer cannot complete in 9 months", "Hire required roles, phase-wise delivery, prioritize core features"),
    ("Delayed Timeline", "Medium", "Features take longer than planned", "Weekly sprint reviews, prioritization, scope management"),
    ("Data Collection Bottleneck", "High", "1000+ colleges data entry is slow", "CSV import, data entry team, template-based entry, outsourcing if needed"),
    ("Natural Language Search", "Medium", "Complex NLP conversion fails", "Start with keyword matching, improve iteratively, use AI for parsing"),
    ("Programmatic SEO Quality", "Medium", "Google penalizes thin content", "Minimum content threshold, noindex thin pages, quality over quantity"),
    ("Infrastructure Cost", "Medium", "Running costs exceed budget", "Start minimal, scale gradually, optimize resource usage"),
    ("Security Breach", "High", "Data leak or hack", "Security from day 1, penetration testing, monitoring, audit logs"),
    ("Third-party API Downtime", "Low", "OpenAI/SendGrid offline", "Fallback providers, retry logic, error handling"),
    ("Scope Creep", "High", "Client requests grow beyond original scope", "Clear scope definition, change request process, milestone sign-offs"),
    ("Server Failure", "Medium", "Production downtime", "Replication, automated backups, monitoring alerts, disaster recovery plan"),
]
add_table(["Risk", "Impact", "Description", "Mitigation"], risks)

doc.add_heading('5.2. Risk Response Strategies', level=2)
add_bullet("AVOID: Clear scope, change management process")
add_bullet("MITIGATE: Regular testing, monitoring, backups")
add_bullet("TRANSFER: Use managed services (cloud, CDN)")
add_bullet("ACCEPT: Low-impact risks with minimal cost to fix")

doc.add_page_break()

# ============================================
# 6. COMMUNICATION PLAN
# ============================================
doc.add_heading('6. COMMUNICATION PLAN', level=1)

doc.add_heading('6.1. Communication Channels', level=2)
comm = [
    ("Daily Standup", "Daily", "Development team", "Progress, blockers, next steps"),
    ("Weekly Status Report", "Weekly (Monday)", "PM to Project Owner", "Progress, milestones, issues, next week plan"),
    ("Sprint Review", "Bi-weekly", "All stakeholders", "Demo completed work, feedback"),
    ("Milestone Review", "End of each month", "Project Owner + PM", "Phase completion, sign-off"),
    ("Change Request", "As needed", "Team + Project Owner", "Scope changes approval"),
    ("Emergency Communication", "Immediate", "PM + Project Owner", "Critical issues, production incidents"),
]
add_table(["Channel", "Frequency", "Participants", "Purpose"], comm)

doc.add_heading('6.2. Reporting Requirements', level=2)
add_bullet("Weekly progress report to Project Owner")
add_bullet("Monthly milestone sign-off document")
add_bullet("Bug/issue tracker updated continuously")
add_bullet("Risk register reviewed weekly")
add_bullet("Budget tracking report monthly")

doc.add_page_break()

# ============================================
# 7. QUALITY MANAGEMENT PLAN
# ============================================
doc.add_heading('7. QUALITY MANAGEMENT PLAN', level=1)

doc.add_heading('7.1. Quality Standards', level=2)
add_bullet("Code follows PEP 8 (Python) and standard JS/TS conventions")
add_bullet("All APIs documented with FastAPI Swagger")
add_bullet("Unit test coverage: minimum 70% for backend")
add_bullet("Performance: API response <200ms for most endpoints")
add_bullet("Accessibility: WCAG 2.1 AA compliance")
add_bullet("Responsive: Works on 320px to 1440px+ screens")

doc.add_heading('7.2. Quality Gates', level=2)
gates = [
    ("Code Review", "Every PR", "Peer reviewer", "No critical issues"),
    ("Unit Tests", "Every build", "CI Pipeline", "All tests pass"),
    ("Integration Tests", "Each phase", "QA Engineer", "Core flows work"),
    ("Performance Test", "Month 8", "QA Engineer", "Response times ok"),
    ("Security Scan", "Month 8-9", "Security specialist", "No critical vulnerabilities"),
    ("UAT", "Month 9", "Project Owner + selected users", "User acceptance signed"),
]
add_table(["Quality Gate", "When", "Responsible", "Pass Criteria"], gates)

doc.add_page_break()

# ============================================
# 8. CHANGE MANAGEMENT
# ============================================
doc.add_heading('8. CHANGE MANAGEMENT PLAN', level=1)

doc.add_heading('8.1. Change Request Process', level=2)
add_bullet("Step 1: Stakeholder submits change request")
add_bullet("Step 2: PM assesses impact (time, cost, scope)")
add_bullet("Step 3: PM presents impact analysis to Project Owner")
add_bullet("Step 4: Project Owner approves/rejects")
add_bullet("Step 5: If approved, update project plan and schedule")
add_bullet("Step 6: Communicate change to affected team members")

doc.add_heading('8.2. Change Categories', level=2)
add_bullet("CRITICAL: Must-have, affects launch (approve immediately)")
add_bullet("IMPORTANT: Should-have, defer if possible")
add_bullet("MINOR: Nice-to-have, add to future phase")

doc.add_page_break()

# ============================================
# 9. DATA COLLECTION & CONTENT PLAN
# ============================================
doc.add_heading('9. DATA COLLECTION & CONTENT PLAN', level=1)

doc.add_heading('9.1. Data Sources', level=2)
add_bullet("Official college websites")
add_bullet("Government education portals")
add_bullet("Regulatory authorities (NMC, AICTE, INC, etc.)")
add_bullet("University websites")
add_bullet("Verified institutional communication")
add_bullet("Public data sources")

doc.add_heading('9.2. Data Collection Strategy', level=2)
content_rows = [
    ("College Data", "500-1000", "Month 2-4", "Data entry team + CSV import", "Official websites, govt sources"),
    ("Course Data", "50-100", "Month 2", "Admin + content team", "Regulatory authorities"),
    ("Scholarship Data", "100-200", "Month 3", "Content team", "Govt scholarships portal"),
    ("Exam Data", "30-50", "Month 3", "Content team", "Official exam websites"),
    ("Mock Test Questions", "2000-5000", "Month 7", "Question creators + import", "Subject experts"),
    ("Blog Articles", "50-100", "Month 6+", "Content writers", "Education research"),
    ("College Images", "Per college", "Ongoing", "Web scraping + submissions", "College websites"),
]
add_table(["Content Type", "Quantity", "Timeline", "Who", "Source"], content_rows)

doc.add_heading('9.3. Data Quality Process', level=2)
add_bullet("Every record has: data_source, verification_status, last_verified_date")
add_bullet("Source must be one of: official website, govt source, regulatory authority, university, verified communication")
add_bullet("Fees marked as 'approximate' when not guaranteed")
add_bullet("Verification team reviews data periodically")
add_bullet("User feedback loop for data corrections")

doc.add_page_break()

# ============================================
# 10. DEPLOYMENT & LAUNCH PLAN
# ============================================
doc.add_heading('10. DEPLOYMENT & LAUNCH PLAN', level=1)

doc.add_heading('10.1. Environment Strategy', level=2)
envs = [
    ("Development", "localhost", "Daily", "Feature development"),
    ("Staging", "staging.padhaanewala.in", "Test releases", "QA, UAT, client demo"),
    ("Production", "padhaanewala.in", "Final", "Live users"),
]
add_table(["Environment", "URL", "Update Frequency", "Purpose"], envs)

doc.add_heading('10.2. Deployment Pipeline', level=2)
add_body("GitHub Actions CI/CD:")
add_bullet("1. Developer pushes code to feature/* branch")
add_bullet("2. CI runs linting + unit tests")
add_bullet("3. Code review + merge to develop branch")
add_bullet("4. Auto-deploy to staging environment")
add_bullet("5. QA testing on staging")
add_bullet("6. Merge to main branch")
add_bullet("7. Auto-deploy to production (after approval)")

doc.add_heading('10.3. Launch Strategy', level=2)
add_bullet("Soft launch with core features (Month 8.5)")
add_bullet("Invite selected beta users for feedback")
add_bullet("Fix critical issues from beta")
add_bullet("Full production launch (Month 9)")
add_bullet("Post-launch monitoring for 48 hours")
add_bullet("Prepare rollback plan if critical issues occur")

doc.add_page_break()

# ============================================
# 11. PROJECT GOVERNANCE
# ============================================
doc.add_heading('11. PROJECT GOVERNANCE', level=1)

doc.add_heading('11.1. Decision Authority', level=2)
governance = [
    ("Daily Decisions", "Development team", "Technical choices within scope"),
    ("Sprint Decisions", "PM", "Task priorities, sprint scope"),
    ("Phase Scope", "PM + Project Owner", "Feature prioritization, scope"),
    ("Budget Decisions", "Project Owner", "Budget allocation, hiring"),
    ("Architecture Changes", "PM + Tech Lead", "Material architecture deviations"),
    ("Final Approval", "Project Owner", "Launch approval, contract"),
]
add_table(["Decision Type", "Decision Maker", "Scope"], governance)

doc.add_heading('11.2. Milestone Sign-off', level=2)
signoffs = [
    ("M1: Foundation Complete", "End of Month 1", "Database schema, auth, environment"),
    ("M2: Core Features Complete", "End of Month 2", "Colleges, Courses, Search"),
    ("M3: Content Complete", "End of Month 3", "Scholarships, Exams, Search"),
    ("M4: Student Features", "End of Month 4", "Accounts, Reviews, Comparison"),
    ("M5: AI Features", "End of Month 5", "Predictor, Assistant"),
    ("M6: Admin Complete", "End of Month 6", "Admin, CMS"),
    ("M7: Mock Tests", "End of Month 7", "Test engine complete"),
    ("M8: CRM Complete", "End of Month 8", "Leads, Counsellors"),
    ("M9: LAUNCH", "End of Month 9", "Production launch"),
]
add_table(["Milestone", "Target Date", "Success Criteria"], signoffs)

doc.add_page_break()

# ============================================
# 12. POST-LAUNCH SUPPORT
# ============================================
doc.add_heading('12. POST-LAUNCH SUPPORT PLAN', level=1)

doc.add_heading('12.1. Support Structure (First 90 Days)', level=2)
support = [
    ("Week 1-2", "24/7 monitoring", "Critical bug fixes within 4 hours"),
    ("Week 3-4", "High priority monitoring", "Bugs fixed within 24 hours"),
    ("Month 2", "Normal priority", "Bugs fixed within 48 hours"),
    ("Month 3", "Maintenance mode", "Scheduled updates"),
]
add_table(["Period", "Monitoring Level", "Response Time"], support)

doc.add_heading('12.2. Maintenance Plan', level=2)
add_bullet("Weekly data verification updates")
add_bullet("Monthly dependency updates")
add_bullet("Quarterly security audits")
add_bullet("Continuous SEO improvements")
add_bullet("Content updates (blogs, scholarships, exams)")

doc.add_heading('12.3. Post-Launch KPIs (First 3 Months)', level=2)
kpis = [
    ("Website visitors", "10,000/month"),
    ("College page views", "5,000/month"),
    ("Course page views", "3,000/month"),
    ("Enquiries/Leads", "500/month"),
    ("Predictor usage", "1,000/month"),
    ("Mock test attempts", "500/month"),
    ("Student registrations", "1,000/month"),
    ("Blog readers", "2,000/month"),
    ("Google ranking (target keywords)", "Top 20"),
]
add_table(["KPI", "Target (3 months)"], kpis)

# ============================================
# SECTION 13: APPROVALS
# ============================================
doc.add_page_break()
doc.add_heading('13. APPROVALS', level=1)

add_body("This Project Plan has been reviewed and approved by:")

doc.add_paragraph()
approvals = [
    ("Padhaanewala Founder / Project Owner", "______________________", "Date: ____________"),
    ("Project Manager", "______________________", "Date: ____________"),
    ("Lead Backend Engineer", "______________________", "Date: ____________"),
    ("Lead Frontend Engineer", "______________________", "Date: ____________"),
]
add_table(["Role", "Signature", "Date"], approvals)

output_path = r"D:\code\Clients\Padhaanewala\Padhaanewala_Project_Plan.docx"
doc.save(output_path)
print(f"Project Plan saved to: {output_path}")
