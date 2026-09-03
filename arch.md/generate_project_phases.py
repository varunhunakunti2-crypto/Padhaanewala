from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.color.rgb = RGBColor(0, 51, 102)
    hs.font.bold = True
    if level == 1:
        hs.font.size = Pt(18)
        hs.paragraph_format.space_before = Pt(24)
        hs.paragraph_format.space_after = Pt(12)
    elif level == 2:
        hs.font.size = Pt(15)
        hs.paragraph_format.space_before = Pt(18)
        hs.paragraph_format.space_after = Pt(8)
    else:
        hs.font.size = Pt(13)
        hs.paragraph_format.space_before = Pt(12)
        hs.paragraph_format.space_after = Pt(6)

def add_title(text, size=24, bold=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0, 51, 102)
    return p

def add_body(text):
    doc.add_paragraph(text)

def add_bullet(text):
    doc.add_paragraph(text, style='List Bullet')

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            table.rows[ri + 1].cells[ci].text = str(val)
    doc.add_paragraph()
    return table

# ============================================
# COVER PAGE
# ============================================
for _ in range(4):
    doc.add_paragraph()

add_title("PADHAANEWALA EDUTECH SERVICES")
add_title("BENGALURU - 560100", size=16, bold=False)
doc.add_paragraph()
add_title("PROJECT PHASES", size=28)
doc.add_paragraph()
add_title("Padhaanewala Education Technology Platform", size=14, bold=False)
add_title("padhaanewala.in", size=12, bold=False)
add_title("Version: 1.0", size=12, bold=False)
add_title("Date: September 2026", size=12, bold=False)

doc.add_page_break()

# ============================================
# PHASE OVERVIEW
# ============================================
doc.add_heading('PROJECT PHASES OVERVIEW', level=1)

add_body("Padhaanewala project follows a phased development approach. Each phase builds on the previous one, ensuring a solid foundation before adding advanced features. The project consists of 9 implementation months across 5 core development phases, categorized as follows:")

doc.add_heading('Two Views of Phases', level=2)
add_body("There are two ways to understand the project phases:")

add_body("1. PROJECT LIFECYCLE PHASES (Management view) - 5 phases that apply to the whole project:")
add_bullet("Phase 0: Initiation")
add_bullet("Phase 1: Planning")
add_bullet("Phase 2: Execution")
add_bullet("Phase 3: Monitoring & Control")
add_bullet("Phase 4: Closing")

add_body("2. DEVELOPMENT BUILD PHASES (Feature view) - 5 phases from the Master Specification, spread over 9 months:")
table_phases = [
    ("PHASE A", "Core Platform", "Month 1-3", "Homepage, College DB, Search, Admin, CMS, Enquiry, Auth, Basic SEO"),
    ("PHASE B", "Student Features", "Month 4", "Dashboard, Saved Colleges, Comparison, Scholarships, Reviews"),
    ("PHASE C", "AI Features", "Month 5", "College Predictor, Education Assistant, AI Comparison (RAG-based)"),
    ("PHASE D", "Mock Tests", "Month 6-7", "Question Bank, Test Engine, Timer, Results, Analytics"),
    ("PHASE E", "Business/CRM", "Month 8-9", "Counsellor Dashboard, Lead Management, CRM, Marketing, Launch"),
]
add_table(["Phase", "Name", "Timeline", "Key Features"], table_phases)

doc.add_page_break()

# ============================================
# PHASE A: CORE PLATFORM
# ============================================
doc.add_heading('PHASE A: CORE PLATFORM (Foundation)', level=1)

doc.add_heading('A.1. Phase Objective', level=2)
add_body("Build the solid foundation of the platform. This phase establishes the technology stack, database, basic infrastructure, and the most critical features. Without this phase, nothing else can function.")

doc.add_heading('A.2. Timeline', level=2)
add_bullet("Duration: Months 1-3 (Weeks 1-12)")
add_bullet("This is the longest phase because it includes infrastructure setup, database design, and core features.")

doc.add_heading('A.3. Components', level=2)

doc.add_heading('1. Project & Infrastructure Setup', level=3)
add_bullet("Next.js + FastAPI + PostgreSQL project initialization")
add_bullet("Docker containerization")
add_bullet("CI/CD pipeline (GitHub Actions)")
add_bullet("Dev, Staging, Production environments")
add_bullet("Environment variables configuration")
add_bullet("Cloudflare CDN + WAF setup")

doc.add_heading('2. Database Design', level=3)
add_bullet("All 45+ database tables created")
add_bullet("Alembic migrations")
add_bullet("pgvector extension for future RAG")
add_bullet("Seed data for development")
add_bullet("Indexes and relationships established")

doc.add_heading('3. Authentication System', level=3)
add_bullet("JWT token-based authentication")
add_bullet("Mobile OTP verification")
add_bullet("Email/Password registration")
add_bullet("Role-based access control (RBAC)")
add_bullet("Password hashing (bcrypt)")
add_bullet("Session management")

doc.add_heading('4. College System', level=3)
add_bullet("College database with all 30+ fields")
add_bullet("College CRUD (Admin API)")
add_bullet("College public pages (/college/slug)")
add_bullet("College image gallery")
add_bullet("CSV/Excel import with validation")
add_bullet("Duplicate detection")
add_bullet("Data verification workflow")

doc.add_heading('5. Course System', level=3)
add_bullet("Course database as separate entities")
add_bullet("Course CRUD (Admin API)")
add_bullet("Course public pages (/courses/slug)")
add_bullet("College-Course relationship (many-to-many)")
add_bullet("Structured fee data")

doc.add_heading('6. College Search', level=3)
add_bullet("Advanced filter system (course, state, district, city, type, fees, hostel, rating)")
add_bullet("Pagination")
add_bullet("Search suggestions/autocomplete")
add_bullet("Natural language search (basic version)")

doc.add_heading('7. Admin Panel (Basic)', level=3)
add_bullet("Dashboard")
add_bullet("College management")
add_bullet("Course management")
add_bullet("User management")
add_bullet("CMS (basic)")

doc.add_heading('8. Enquiry System', level=3)
add_bullet("Enquiry forms on all important pages")
add_bullet("Lead capture in CRM")
add_bullet("Source tracking")

doc.add_heading('9. Basic SEO', level=3)
add_bullet("Clean URLs")
add_bullet("Meta titles and descriptions")
add_bullet("Sitemap")
add_bullet("Robots.txt")
add_bullet("Schema.org markup basics")

doc.add_heading('A.4. Deliverables', level=2)
add_bullet("Working development environment")
add_bullet("Complete database schema")
add_bullet("College and course management system")
add_bullet("Working search functionality")
add_bullet("Basic admin panel")
add_bullet("Authentication system")
add_bullet("Enquiry capture")
add_bullet("SEO foundation")

doc.add_heading('A.5. Exit Criteria', level=2)
add_bullet("All core APIs tested and working")
add_bullet("College and course pages render correctly")
add_bullet("Admin can add/edit/delete colleges and courses")
add_bullet("Enquiries reach the admin panel")
add_bullet("Authentication works for admin and public")
add_bullet("Database backup working")

doc.add_page_break()

# ============================================
# PHASE B: STUDENT FEATURES
# ============================================
doc.add_heading('PHASE B: STUDENT FEATURES', level=1)

doc.add_heading('B.1. Phase Objective', level=2)
add_body("Add student-facing features that allow users to create accounts, save colleges, compare institutions, submit reviews, and manage their education journey. This phase transforms the platform from a content site into an interactive student experience.")

doc.add_heading('B.2. Timeline', level=2)
add_bullet("Duration: Month 4 (Weeks 13-16)")

doc.add_heading('B.3. Components', level=2)

doc.add_heading('1. Student Registration & Login', level=3)
add_bullet("Mobile OTP registration")
add_bullet("Email/Password registration")
add_bullet("Social login (optional)")
add_bullet("Password reset")
add_bullet("Profile completion flow")

doc.add_heading('2. Student Dashboard', level=3)
add_bullet("Personal profile management")
add_bullet("Education details")
add_bullet("Course interest")
add_bullet("Preferred states/cities")
add_bullet("Budget setting")
add_bullet("Activity overview")

doc.add_heading('3. Saved Colleges', level=3)
add_bullet("Save College button on college pages")
add_bullet("My Colleges section in dashboard")
add_bullet("Remove saved college")
add_bullet("Compare saved colleges")

doc.add_heading('4. College Comparison', level=3)
add_bullet("Select multiple colleges (2-4)")
add_bullet("Side-by-side comparison view")
add_bullet("Compare: Location, Type, University, Course, Duration, Fees, Hostel, Facilities, Admission, Reviews")
add_bullet("\"Ask AI: Which college is better for me?\" feature (Phase C)")

doc.add_heading('5. Scholarship Finder', level=3)
add_bullet("Scholarship database (100-200 records)")
add_bullet("Scholarship filters (course, state, income, deadline)")
add_bullet("Scholarship detail pages")
add_bullet("Official application links (clearly marked)")

doc.add_heading('6. Exam Notifications', level=3)
add_bullet("Exam database (30-50 records)")
add_bullet("Exam detail pages")
add_bullet("Exam dates display")
add_bullet("Deadline tracking")

doc.add_heading('7. Review System', level=3)
add_bullet("Students submit college reviews")
add_bullet("Rating + review text + optional images")
add_bullet("Moderation workflow: Submitted → Moderation → Approved → Published")
add_bullet("Admin can approve/reject/spam-filter reviews")

doc.add_heading('B.4. Deliverables', level=2)
add_bullet("Student accounts and profiles")
add_bullet("Saved colleges functionality")
add_bullet("College comparison tool")
add_bullet("Scholarship finder with filters")
add_bullet("Exam information pages")
add_bullet("Review submission and moderation system")

doc.add_heading('B.5. Exit Criteria', level=2)
add_bullet("Students can register and log in securely")
add_bullet("Saved colleges appear in dashboard")
add_bullet("Comparison tool works with 2-4 colleges")
add_bullet("Scholarship search returns filtered results")
add_bullet("Reviews go through moderation before publishing")

doc.add_page_break()

# ============================================
# PHASE C: AI FEATURES
# ============================================
doc.add_heading('PHASE C: AI FEATURES (RAG-based)', level=1)

doc.add_heading('C.1. Phase Objective', level=2)
add_body("Implement the AI-powered features that make Padhaanewala stand out. All AI features use RAG (Retrieval-Augmented Generation) to ensure responses are based on verified database information, not AI hallucinations.")

doc.add_heading('C.2. Timeline', level=2)
add_bullet("Duration: Month 5 (Weeks 17-20)")
add_bullet("Requires Phase A (database with data) and Phase B (student features) to be complete")

doc.add_heading('C.3. RAG Infrastructure (Prerequisite)', level=2)

doc.add_heading('1. pgvector Setup', level=3)
add_bullet("Install pgvector extension")
add_bullet("Create document_embeddings table")
add_bullet("Configure vector indexes (IVFFlat)")

doc.add_heading('2. Embedding Generation', level=3)
add_bullet("Embedding service (OpenAI/Cohere)")
add_bullet("Chunking strategy for colleges, courses, scholarships, FAQs")
add_bullet("Celery task for embedding on data changes")
add_bullet("Embedding storage and management")

doc.add_heading('3. Vector Search', level=3)
add_bullet("Cosine similarity search")
add_bullet("Top-K retrieval (5-10 chunks)")
add_bullet("Hybrid search (DB filters + vector similarity)")

doc.add_heading('C.4. AI College Predictor', level=2)
add_body("One of the main Padhaanewala features.")
add_bullet("Inputs: Course, Entrance Exam, Rank/Score, Category, State, City, Budget, College Type, Hostel")
add_bullet("RAG Flow:")
add_bullet("  1. Extract structured filters from student inputs")
add_bullet("  2. Query PostgreSQL for eligible colleges")
add_bullet("  3. Generate embedding for preferences")
add_bullet("  4. Vector search for similar colleges")
add_bullet("  5. Combine results and rank using AI")
add_bullet("  6. Categorize: Highly Suitable / Possible / Reach")
add_bullet("  7. Generate explanation from verified data")
add_bullet("  8. Add disclaimer: \"Results are estimates, not guaranteed admissions\"")

doc.add_heading('C.5. AI Education Assistant', level=2)
add_body("\"Ask Padhaanewala AI\" chatbot.")
add_bullet("Questions: What is BHMS? Difference between BAMS and BHMS? Which course after 12th?")
add_bullet("RAG Flow:")
add_bullet("  1. Preprocess question (intent + entities)")
add_bullet("  2. Search PostgreSQL for course/college data")
add_bullet("  3. Vector search for similar content")
add_bullet("  4. Retrieve FAQs")
add_bullet("  5. Assemble context")
add_bullet("  6. Generate answer with LLM")
add_bullet("  7. Cite sources")
add_bullet("  8. Add verification disclaimer")

doc.add_heading('C.6. AI College Comparison', level=2)
add_bullet("\"Ask AI: Which college is better for me?\"")
add_bullet("Uses verified data from both colleges")
add_bullet("AI explains the comparison with reasons")
add_bullet("Clearly states recommendations are not admission guarantees")

doc.add_heading('C.7. Natural Language Search (Enhanced)', level=2)
add_bullet("Parse: \"Nursing colleges near Bangalore under 5 lakh\"")
add_bullet("Extract: Course, Location, Max Fee")
add_bullet("Execute structured search + vector search")
add_bullet("Return explanations with results")

doc.add_heading('C.8. AI Guardrails', level=2)
add_bullet("No hallucination - AI only uses retrieved database info")
add_bullet("Source citation - every AI response cites data sources")
add_bullet("Uncertainty disclosure - AI states when info is uncertain")
add_bullet("Verification reminder - always tell users to verify with institutions")
add_bullet("No guarantees - never promise admissions")

doc.add_heading('C.9. Deliverables', level=2)
add_bullet("AI College Predictor (RAG-based)")
add_bullet("AI Education Assistant (RAG-based)")
add_bullet("AI College Comparison")
add_bullet("Enhanced Natural Language Search")
add_bullet("document_embeddings table with data")
add_bullet("Vector search API")

doc.add_heading('C.10. Exit Criteria', level=2)
add_bullet("AI College Predictor returns categorized results with explanations")
add_bullet("AI Education Assistant answers questions from verified data")
add_bullet("All AI responses cite sources")
add_bullet("No AI hallucination observed in testing (AI only uses DB data)")
add_bullet("Disclaimers present on all AI outputs")

doc.add_page_break()

# ============================================
# PHASE D: MOCK TESTS
# ============================================
doc.add_heading('PHASE D: MOCK TESTS', level=1)

doc.add_heading('D.1. Phase Objective', level=2)
add_body("Build a comprehensive mock test system that allows students to practice for entrance exams. Includes question bank, test engine, timer, scoring, and performance analytics.")

doc.add_heading('D.2. Timeline', level=2)
add_bullet("Duration: Months 6-7 (Weeks 25-28)")

doc.add_heading('D.3. Components', level=2)

doc.add_heading('1. Question Bank', level=3)
add_bullet("Question CRUD (Admin)")
add_bullet("Question import (CSV)")
add_bullet("Question categories/difficulty levels")
add_bullet("Question validation")
add_bullet("Subject-wise organization")
add_bullet("Multiple question types (MCQ, etc.)")
add_bullet("Explanation/answer key for each question")

doc.add_heading('2. Test Creation (Admin)', level=3)
add_bullet("Create mock tests by exam")
add_bullet("Select subjects")
add_bullet("Set difficulty level")
add_bullet("Set number of questions")
add_bullet("Set time limit")
add_bullet("Publish/unpublish tests")

doc.add_heading('3. Test Taking Interface (Student)', level=3)
add_bullet("Exam selection")
add_bullet("Subject selection")
add_bullet("Difficulty selection")
add_bullet("Timer display")
add_bullet("Question navigation (Next/Previous)")
add_bullet("Mark for review")
add_bullet("Question palette (answered, not answered, marked)")
add_bullet("Submit test")

doc.add_heading('4. Test Results & Analytics', level=3)
add_bullet("Score")
add_bullet("Percentage")
add_bullet("Correct answers count")
add_bullet("Incorrect answers count")
add_bullet("Unattempted questions count")
add_bullet("Time taken")
add_bullet("Topic-wise performance breakdown")
add_bullet("Rank/percentile (where meaningful)")
add_bullet("Practice Again button")
add_bullet("View Solutions button")

doc.add_heading('5. Test History', level=3)
add_bullet("All past attempts visible in student dashboard")
add_bullet("Performance trends over time")
add_bullet("Improvement tracking")

doc.add_heading('D.4. Deliverables', level=2)
add_bullet("Question bank (target 2000-5000 questions)")
add_bullet("Test creation tool (Admin)")
add_bullet("Test taking interface (Student)")
add_bullet("Scoring and results system")
add_bullet("Performance analytics dashboard")

doc.add_heading('D.5. Exit Criteria', level=2)
add_bullet("Admin can create tests with questions")
add_bullet("Students can take tests with timer")
add_bullet("Results calculate correctly (score, percentage, correct/incorrect)")
add_bullet("Topic-wise performance displayed")
add_bullet("Test history visible in dashboard")

doc.add_page_break()

# ============================================
# PHASE E: BUSINESS/CRM
# ============================================
doc.add_heading('PHASE E: BUSINESS/CRM & LAUNCH', level=1)

doc.add_heading('E.1. Phase Objective', level=2)
add_body("Add the business-facing features that turn Padhaanewala into a lead-generation machine. This includes counsellor dashboards, lead management, CRM, marketing attribution, and takes the platform to production launch.")

doc.add_heading('E.2. Timeline', level=2)
add_bullet("Duration: Months 8-9 (Weeks 29-36)")

doc.add_heading('E.3. Components', level=2)

doc.add_heading('1. Lead Management', level=3)
add_bullet("Lead ID generation")
add_bullet("Lead fields: Name, Mobile, Email, Course, College, Source, Date, Status, Assigned counsellor, Follow-up date, Notes")
add_bullet("Lead statuses: New, Contacted, Interested, Application Started, Admission Completed, Not Interested, Closed")
add_bullet("Lead list with filters and search")
add_bullet("Lead detail view")

doc.add_heading('2. Counsellor Dashboard', level=3)
add_bullet("Counsellor login")
add_bullet("Assigned leads list")
add_bullet("Student details view")
add_bullet("Add/update notes")
add_bullet("Schedule follow-ups")
add_bullet("Update contact status")
add_bullet("Update admission status")
add_bullet("Role-based access (counsellors only see their own leads)")

doc.add_heading('3. Lead Assignment', level=3)
add_bullet("Admin assigns leads to counsellors")
add_bullet("Auto-assignment rules (optional)")
add_bullet("Reassignment capability")
add_bullet("Workload balancing")

doc.add_heading('4. Lead Source Tracking', level=3)
add_bullet("Source tracking: Homepage, College page, Course page, Blog, Scholarship, Predictor, WhatsApp, Advertisement, Organic search")
add_bullet("UTM tracking (utm_source, utm_medium, utm_campaign, utm_content)")
add_bullet("Marketing attribution reports")

doc.add_heading('5. CRM Features', level=3)
add_bullet("Lead pipeline view")
add_bullet("Follow-up reminders")
add_bullet("Activity history per lead")
add_bullet("Communication log")
add_bullet("Conversion tracking")

doc.add_heading('6. Analytics Dashboard (Enhanced)', level=3)
add_bullet("Website: Visitors, Searches, College views, Course views, Predictor usage")
add_bullet("Students: Registrations, Active users")
add_bullet("Leads: New, Contacted, Converted")
add_bullet("Content: Most viewed colleges, Most searched courses, Popular scholarships, Popular articles")

doc.add_heading('7. Marketing Attribution', level=3)
add_bullet("Compare effectiveness of lead sources")
add_bullet("Conversion rate by source")
add_bullet("ROI tracking by campaign")

doc.add_heading('8. Final Production Launch', level=3)
add_bullet("Production server setup")
add_bullet("Domain/DNS configuration")
add_bullet("SSL certificate")
add_bullet("CDN setup")
add_bullet("Monitoring setup")
add_bullet("Backup automation")
add_bullet("Load testing")
add_bullet("Security testing")
add_bullet("Final QA")
add_bullet("Production launch")

doc.add_heading('E.4. Deliverables', level=2)
add_bullet("Complete lead management system")
add_bullet("Counsellor dashboard with role-based access")
add_bullet("Follow-up and notes system")
add_bullet("Marketing attribution reports")
add_bullet("Production launch")
add_bullet("Post-launch monitoring")

doc.add_heading('E.5. Exit Criteria', level=2)
add_bullet("Leads reach admin panel immediately after enquiry")
add_bullet("Counsellors can manage assigned leads")
add_bullet("Lead sources tracked correctly")
add_bullet("Analytics dashboard functional")
add_bullet("Platform live in production")
add_bullet("All launch checklist items complete")

doc.add_page_break()

# ============================================
# PHASES COMPARISON TABLE
# ============================================
doc.add_heading('PHASES COMPARISON SUMMARY', level=1)

doc.add_heading('All 5 Development Phases at a Glance', level=2)
summary = [
    ("A", "Core Platform", "1-3", "Month 1-3", "Foundation, Database, Colleges, Courses, Search, Admin, Auth, SEO", "The base - nothing else works without this"),
    ("B", "Student Features", "4", "Month 4", "Accounts, Saved Colleges, Comparison, Scholarships, Reviews", "Turns site into interactive experience"),
    ("C", "AI Features", "5", "Month 5", "Predictor, Assistant, AI Comparison (RAG)", "Unique competitive advantage"),
    ("D", "Mock Tests", "6-7", "Month 6-7", "Question Bank, Test Engine, Results, Analytics", "Engagement + recurring usage"),
    ("E", "Business/CRM", "8-9", "Month 8-9", "Leads, Counsellors, CRM, Attribution, Launch", "Revenue generation + production"),
]
add_table(["Phase", "Name", "Weeks", "Months", "Key Features", "Purpose"], summary)

doc.add_page_break()

# ============================================
# PROJECT LIFECYCLE PHASES
# ============================================
doc.add_heading('PROJECT LIFECYCLE PHASES (Management View)', level=1)

add_body("In addition to the development build phases, the project follows a standard project management lifecycle with 5 phases. These apply to the entire project from start to finish.")

doc.add_heading('Phase L1: INITIATION', level=2)
add_bullet("Goal: Define the project, get approval, establish feasibility")
add_bullet("Activities:")
add_bullet("  - Review and approve Master Specification")
add_bullet("  - Define project scope and objectives")
add_bullet("  - Identify stakeholders")
add_bullet("  - Establish project charter")
add_bullet("  - Determine budget and resources")
add_bullet("  - Assess feasibility")
add_bullet("Deliverable: Approved Project Charter, Master Specification (Version 2.0)")
add_bullet("Status: COMPLETE ✅")

doc.add_heading('Phase L2: PLANNING', level=2)
add_bullet("Goal: Create a detailed roadmap for execution")
add_bullet("Activities:")
add_bullet("  - Create detailed project plan")
add_bullet("  - Design system architecture")
add_bullet("  - Design database schema")
add_bullet("  - Create API design")
add_bullet("  - Define team roles and responsibilities")
add_bullet("  - Prepare risk management plan")
add_bullet("  - Define quality standards")
add_bullet("  - Set budget and timeline")
add_bullet("Deliverable: Project Plan, Architecture Design, Database Schema, API Design")
add_bullet("Status: IN PROGRESS (This document) 📋")

doc.add_heading('Phase L3: EXECUTION', level=2)
add_bullet("Goal: Build the platform according to the plan")
add_bullet("Activities:")
add_bullet("  - Set up development environment")
add_bullet("  - Build database and backend")
add_bullet("  - Build frontend")
add_bullet("  - Implement features phase by phase")
add_bullet("  - Integrate AI/RAG features")
add_bullet("  - Update documentation")
add_bullet("Deliverable: Working software (Phase A → B → C → D → E)")
add_bullet("Status: NOT STARTED ⏳")

doc.add_heading('Phase L4: MONITORING & CONTROL', level=2)
add_bullet("Goal: Track progress, ensure quality, fix issues")
add_bullet("Activities:")
add_bullet("  - Track progress against plan")
add_bullet("  - Run tests (unit, integration, E2E)")
add_bullet("  - Perform quality assurance")
add_bullet("  - Conduct security testing")
add_bullet("  - Monitor performance")
add_bullet("  - Review budget and timeline")
add_bullet("  - Manage risks and changes")
add_bullet("Deliverable: Test Reports, Status Reports, Quality Certifications")
add_bullet("Status: Runs continuously throughout execution")

doc.add_heading('Phase L5: CLOSING', level=2)
add_bullet("Goal: Deliver final product, handover, close project")
add_bullet("Activities:")
add_bullet("  - Production deployment")
add_bullet("  - Final testing and verification")
add_bullet("  - Documentation handover")
add_bullet("  - Admin training")
add_bullet("  - Client acceptance sign-off")
add_bullet("  - Lessons learned review")
add_bullet("  - Post-launch support plan")
add_bullet("Deliverable: Live Platform, Documentation, Handover, Training")
add_bullet("Status: NOT STARTED ⏳")

doc.add_page_break()

# ============================================
# GANTT-STYLE SUMMARY
# ============================================
doc.add_heading('PROJECT TIMELINE (Gantt-style Summary)', level=1)

add_body("Visual representation of months vs phases:")

gantt = [
    ("Phase A: Core Platform", "M1", "Y", "Y", "Y", "-", "-", "-", "-", "-"),
    ("Phase B: Student Features", "M2", "-", "-", "-", "Y", "-", "-", "-", "-"),
    ("Phase C: AI Features", "M3", "-", "-", "-", "-", "Y", "-", "-", "-"),
    ("Phase D: Mock Tests", "M4-5", "-", "-", "-", "-", "-", "Y", "Y", "-"),
    ("Phase E: CRM + Launch", "M6-7", "-", "-", "-", "-", "-", "-", "-", "Y"),
    ("Continuous: Testing/QA", "All", "Y", "Y", "Y", "Y", "Y", "Y", "Y", "Y"),
    ("Continuous: Data Collection", "M2+", "-", "Y", "Y", "Y", "Y", "Y", "Y", "Y"),
]
header = ["Activity", "Duration", "M3", "M4", "M5", "M6", "M7", "M8", "M9", "M10"]
# Note: adjust headers to match 9 months - simplify
gantt_table = [
    ("Phase A: Core Platform", "Months 1-3", "Weeks 1-12", "Homepage, DB, Colleges, Courses, Search, Admin, Auth, SEO"),
    ("Phase B: Student Features", "Month 4", "Weeks 13-16", "Accounts, Saved, Compare, Scholarships, Reviews"),
    ("Phase C: AI Features", "Month 5", "Weeks 17-20", "Predictor, Assistant, AI Comparison"),
    ("Phase D: Mock Tests", "Months 6-7", "Weeks 25-28", "Question Bank, Test Engine, Analytics"),
    ("Phase E: CRM + Launch", "Months 8-9", "Weeks 29-36", "Leads, Counsellors, CRM, Production Launch"),
    ("Testing & QA", "Continuous", "All months", "Unit, Integration, E2E, Security, Performance"),
    ("Data Collection", "Month 2+", "Ongoing", "Colleges, Courses, Scholarships, Questions"),
]
add_table(["Activity", "Duration", "Weeks", "Details"], gantt_table)

doc.add_paragraph()
add_title("PADHAANEWALA EDUTECH SERVICES", size=14)
add_title("BENGALURU - 560100", size=12, bold=False)
add_title("End of Project Phases Document", size=12, bold=False)

output_path = r"D:\code\Clients\Padhaanewala\Padhaanewala_Project_Phases.docx"
doc.save(output_path)
print(f"Project Phases saved to: {output_path}")
