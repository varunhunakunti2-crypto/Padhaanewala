from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
import os

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.color.rgb = RGBColor(0, 51, 102)
    hs.font.bold = True
    if level == 1:
        hs.font.size = Pt(20)
        hs.paragraph_format.space_before = Pt(24)
        hs.paragraph_format.space_after = Pt(12)
    elif level == 2:
        hs.font.size = Pt(16)
        hs.paragraph_format.space_before = Pt(18)
        hs.paragraph_format.space_after = Pt(8)
    else:
        hs.font.size = Pt(13)
        hs.paragraph_format.space_before = Pt(12)
        hs.paragraph_format.space_after = Pt(6)

def add_title(text, size=28, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER):
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0, 51, 102)
    return p

def add_body(text):
    p = doc.add_paragraph(text)
    return p

def add_bullet(text):
    p = doc.add_paragraph(text, style='List Bullet')
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            table.rows[ri + 1].cells[ci].text = str(val)
    doc.add_paragraph()
    return table

# ============================================
# COVER PAGE
# ============================================
for _ in range(4):
    doc.add_paragraph()

add_title("PADHAANEWALA EDUTECH SERVICES", size=24)
add_title("BENGALURU – 560100", size=14, bold=False)
doc.add_paragraph()
add_title("MASTER WEBSITE DEVELOPMENT SPECIFICATION", size=28)
doc.add_paragraph()
add_title("Version: 2.0 (Enhanced with RAG + Solutions)", size=14, bold=False)
add_title("Date: September 2026", size=12, bold=False)
doc.add_paragraph()
doc.add_paragraph()
add_title("Assigned Role: Website Engineering & Product Development Team", size=12, bold=False)
add_title("Project: Padhaanewala Education Technology Platform", size=12, bold=False)
add_title("Website: padhaanewala.in", size=12, bold=False)
add_title("Document: Complete Product, Technical & Development Requirements", size=12, bold=False)

doc.add_page_break()

# ============================================
# SECTION A: ORIGINAL SPECIFICATION
# ============================================
doc.add_heading('SECTION A: ORIGINAL SPECIFICATION', level=1)

sections = [
    ("1. PROJECT OBJECTIVE", [
        "Develop Padhaanewala as a professional, scalable, database-driven education platform.",
        "The website must NOT be developed as a collection of static pages.",
        "It must be a complete platform consisting of: Student-facing website; College database; Course database; Scholarship database; Examination database; Mock-test system; AI College Predictor; AI Education Assistant; College comparison; Student accounts; Reviews; Admission enquiry system; Counsellor/lead management; Admin dashboard; Content Management System; SEO infrastructure; Analytics.",
        "The architecture must allow the platform to expand to a large number of colleges, courses and users without requiring a complete rewrite."
    ]),
    ("2. CORE BUSINESS PURPOSE", [
        "Padhaanewala should help students: discover colleges; search courses; compare colleges; understand eligibility; check fees; understand admission procedures; find scholarships; find examination information; take mock tests; use an AI college predictor; ask education-related questions; save colleges; submit admission enquiries; receive counselling assistance.",
        "The platform should simultaneously generate and manage admission leads for Padhaanewala."
    ]),
    ("3. TECHNOLOGY STACK", [
        "Frontend: Next.js, React, TypeScript, Tailwind CSS, responsive/mobile-first design. Next.js is preferred over a basic React SPA because SEO is extremely important for college/course pages.",
        "Backend: Python, FastAPI, REST API, API versioning.",
        "Database: PostgreSQL.",
        "Cache: Redis where required.",
        "Search: PostgreSQL search initially; Elasticsearch/OpenSearch when scale requires it.",
        "Storage: AWS S3, Cloudflare R2 or equivalent S3-compatible object storage. Do not store large images directly in PostgreSQL.",
        "AI: AI requests must go through the backend. Never expose AI API keys in frontend code."
    ]),
    ("4. DOMAIN AND DEVELOPMENT ENVIRONMENT", [
        "Development may use localhost and staging before the production domain is configured.",
        "Production domain: padhaanewala.in",
        "Maintain separate Development, Staging and Production environments.",
        "Never test major experimental changes directly on production."
    ]),
    ("5. WEBSITE NAVIGATION", [
        "Main navigation: Home, Colleges, Courses, College Predictor, Scholarships, Mock Tests, Exams, Reviews, Blog/Resources, About, Contact.",
        "Also provide Login, Register and Get Admission Help."
    ]),
    ("6. HOMEPAGE", [
        "The homepage must look like a modern education technology platform and NOT like a generic coaching-centre website.",
        "Hero: \"Find the Right College for Your Future\".",
        "Provide a large search bar: \"Search colleges, courses, exams or locations\". Examples: BHMS, BAMS, BUMS, MBBS, BDS, B.Sc Nursing, B.Pharm, D.Pharm, BCA, MBA, Engineering.",
        "Buttons: Search; AI College Predictor.",
        "Quick-action cards: Find Colleges; Compare Colleges; College Predictor; Scholarships; Mock Tests; Admission Assistance.",
        "Homepage sections: Popular courses; Featured colleges; Popular college searches; Scholarships; Upcoming examinations; Mock tests; Why Padhaanewala; Student reviews; Latest education articles; Admission assistance CTA.",
        "All homepage content must be manageable from the admin panel."
    ]),
    ("7. COLLEGE DATABASE", [
        "Create a structured college database. Every college must have a unique ID, e.g. COLLEGE000001.",
        "Fields: College ID; College name; Official name; College type; Government/private; University; State; District; City; Address; Pincode; Website; Email; Phone; Established year; Accreditation; Recognition; Courses; Fees; Hostel; Facilities; Admission information; Eligibility; Entrance examination; Cutoff; Images; Gallery; Reviews; FAQs; Data source; Last verified date; Verification status."
    ]),
    ("8. COLLEGE PAGE", [
        "Every college must have an individual SEO-friendly page, e.g. /college/college-name.",
        "Header: College name; Location; College type; Rating/reviews; Apply/Get Admission Help; Compare; Save College.",
        "Sections: Overview; Courses; Fees; Eligibility; Admission; Cutoff; Facilities; Reviews; Gallery; FAQs."
    ]),
    ("9. COURSE DATABASE", [
        "Courses must be separate database entities.",
        "Examples: MBBS; BDS; BAMS; BHMS; BUMS; B.Sc Nursing; D.Pharm; B.Pharm; BCA; BBA; MBA; Engineering; Paramedical courses; Skill courses.",
        "Each course should have: Name; Degree; Duration; Eligibility; Entrance examination; Admission procedure; Fees information; Colleges offering the course; Career information; FAQs; SEO metadata."
    ]),
    ("10. COURSE PAGE", [
        "Example: /courses/bhms",
        "Sections: Overview; Duration; Eligibility; Admission; Entrance exam; Fees; Colleges; Career opportunities; FAQs; Related courses.",
        "CTA: \"Need help choosing a college?\" → Get Admission Assistance."
    ]),
    ("11. COLLEGE SEARCH", [
        "Create advanced college search.",
        "Filters: Course; State; District; City; Government/private; University; Fees; Hostel; Rating; Accreditation; Admission status.",
        "Example: BHMS + Karnataka + Private returns matching colleges."
    ]),
    ("12. NATURAL LANGUAGE SEARCH", [
        "Allow users to search naturally, for example: \"BHMS colleges in Karnataka\"; \"BAMS colleges near Bangalore\"; \"Nursing colleges under 5 lakh\"; \"Private BHMS colleges with hostel\".",
        "The backend should convert natural-language requests into appropriate structured search filters wherever possible."
    ]),
    ("13. COLLEGE COMPARISON", [
        "Students should be able to select multiple colleges.",
        "Comparison should show: Location; Type; University; Course; Duration; Fees; Hostel; Facilities; Admission; Reviews; other relevant information.",
        "Include: \"Ask AI: Which college is better for me?\" AI recommendations must be based on available verified database information and clearly state that recommendations are not admission guarantees."
    ]),
    ("14. AI COLLEGE PREDICTOR", [
        "This should be one of the main Padhaanewala features.",
        "Inputs: Course; Entrance exam; Rank/score; Category where relevant; State; Preferred city; Budget; Government/private preference; Hostel requirement; Other preferences.",
        "Output: Highly Suitable; Possible; Reach.",
        "The predictor must clearly state that results are estimates and not guaranteed admissions."
    ]),
    ("15. AI ARCHITECTURE", [
        "Recommended flow: Student → Frontend → FastAPI → Database/Search Engine → Relevant verified data → AI processing → Response.",
        "Do NOT allow the AI model to independently invent factual college information.",
        "The AI should use database information for factual information such as College names, Courses, Fees, Locations, Eligibility and Admission information.",
        "Where current information is uncertain, the system should indicate that the information needs verification."
    ]),
    ("16. AI EDUCATION ASSISTANT", [
        "Add: \"Ask Padhaanewala AI\".",
        "Possible questions: What is BHMS? What is the difference between BAMS and BHMS? Which course is suitable after 12th? Which colleges offer B.Sc Nursing? What scholarships are available? How does admission work?",
        "The system should not present changing regulatory/admission information as permanently fixed facts."
    ]),
    ("17. SCHOLARSHIP FINDER", [
        "Create a separate scholarship database.",
        "Fields: Scholarship name; Provider; Government/private; Eligibility; State; Course; Income criteria; Amount; Deadline; Documents; Application procedure; Official application source; Status; Last verified date.",
        "Filters: Course; State; Student category where applicable; Income; Government/private; Deadline."
    ]),
    ("18. SCHOLARSHIP PAGE", [
        "Each scholarship should have: Scholarship name; Provider; Amount; Eligibility; Deadline; Required documents; Application process; Official application link; FAQs.",
        "Clearly distinguish Official scholarship application from Padhaanewala counselling/admission assistance."
    ]),
    ("19. EXAMINATION DATABASE", [
        "Create /exams.",
        "Each exam should include: Exam name; Conducting authority; Eligibility; Application start date; Application deadline; Exam date; Admit card date; Result date; Official website; Official notification; FAQs.",
        "All dates must be editable through the admin panel."
    ]),
    ("20. MOCK TEST SYSTEM", [
        "Create /mock-tests.",
        "Features: Exam selection; Subject selection; Difficulty; Number of questions; Timer; Question navigation; Mark for review; Next/previous; Submit."
    ]),
    ("21. MOCK TEST RESULT", [
        "After submission show: Score; Percentage; Correct answers; Incorrect answers; Unattempted; Time taken; Topic-wise performance; Rank/percentile where meaningful.",
        "Buttons: Practice Again; View Solutions."
    ]),
    ("22. STUDENT REGISTRATION", [
        "Students should be able to register using Mobile OTP and/or Email/Password.",
        "Do not store passwords as plain text. Use secure password hashing."
    ]),
    ("23. STUDENT PROFILE", [
        "Profile should contain: Name; Mobile; Email; Education; Course interest; Preferred state; Preferred city; Budget; Saved colleges; Test history; Scholarship interests; Enquiries."
    ]),
    ("24. SAVE COLLEGE", [
        "Students should be able to click Save College.",
        "Saved colleges should appear under My Colleges.",
        "Students should then be able to compare them."
    ]),
    ("25. ADMISSION ENQUIRY", [
        "Every important page should have Get Admission Assistance.",
        "Form: Name; Mobile; Email; Course; Preferred college; State; City; Qualification; Message.",
        "After submission: \"Thank you. Our counsellor will contact you.\"",
        "The enquiry must immediately enter the admin/CRM system."
    ]),
    ("26. LEAD MANAGEMENT", [
        "Each lead should have: Lead ID; Student name; Mobile; Email; Course; College; Source; Date; Status; Assigned counsellor; Follow-up date; Notes.",
        "Statuses: New; Contacted; Interested; Application Started; Admission Completed; Not Interested; Closed."
    ]),
    ("27. COUNSELLOR DASHBOARD", [
        "Counsellors should have their own login.",
        "Features: Assigned leads; Student details; Notes; Follow-up; Contact status; Admission status.",
        "Use role-based access so counsellors cannot access information they are not authorized to see."
    ]),
    ("28. ADMIN PANEL", [
        "Admin must be able to manage the entire platform.",
        "Modules: Dashboard; Colleges; Courses; Scholarships; Exams; Mock Tests; Questions; Students; Reviews; Blogs; FAQs; Banners; Notifications; Leads; Counsellors; Media; SEO; Settings; Audit logs."
    ]),
    ("29. ADMIN ROLES", [
        "Super Admin: Full access.",
        "Content Admin: College/course/blog/scholarship content.",
        "Counsellor: Assigned leads.",
        "Test Admin: Mock tests/questions.",
        "SEO Admin: SEO metadata/content.",
        "Permissions must be configurable."
    ]),
    ("30. CMS", [
        "Non-technical administrators must be able to update website content.",
        "Admin should be able to Add, Edit, Delete, Publish, Unpublish and Schedule content.",
        "This should include Colleges, Courses, Scholarships, Exams, Blogs, FAQs, Homepage content and Banners."
    ]),
    ("31. BLOG", [
        "Create /blog.",
        "Categories: Admissions; NEET; AYUSH; Nursing; Scholarships; Careers; Exams; College guides; Education news.",
        "Each article needs: Title; Slug; Content; Featured image; Category; Author; Meta title; Meta description; Canonical URL; Publish status."
    ]),
    ("32. SEO", [
        "Every important page must have: SEO title; Meta description; Canonical URL; Open Graph metadata; Structured data; Sitemap; Robots.txt; Clean URL.",
        "Example: /colleges/bhms-colleges-in-karnataka is preferable to /page?id=123."
    ]),
    ("33. PROGRAMMATIC SEO", [
        "The platform may automatically create useful pages such as BHMS colleges in Karnataka, BAMS colleges in Karnataka, Nursing colleges in Bihar and B.Pharm colleges in Bangalore.",
        "Do NOT generate thousands of low-quality duplicate pages. Only useful, sufficiently informative pages should be indexable."
    ]),
    ("34. COLLEGE DATA VERIFICATION", [
        "Every important college record should contain Data source; Verification status; Last verified date; Verified by.",
        "Possible sources: Official college website; Official government source; Regulatory authority; University; Verified institutional communication."
    ]),
    ("35. FEES DATA", [
        "Fees should be structured rather than a single text field.",
        "Possible fields: Tuition fee; Hostel fee; Examination fee; Other charges; Total approximate fee; Fee period.",
        "Display \"Approximate fee\" when exact/current fee cannot be guaranteed.",
        "Include: \"Fees should be verified with the institution before admission.\""
    ]),
    ("36. REVIEWS", [
        "Students can submit: College; Course; Year; Rating; Review; Optional images.",
        "Reviews must go through moderation: Submitted → Moderation → Approved → Published.",
        "Admin should be able to reject spam, abusive or inappropriate content."
    ]),
    ("37. COLLEGE GALLERY", [
        "Images should be stored using object storage.",
        "Database should store: Image URL; College ID; Image type; Alt text; Upload date.",
        "Images must be optimized before delivery."
    ]),
    ("38. WHATSAPP", [
        "Add WhatsApp CTA where appropriate.",
        "Example contextual message: \"Hello Padhaanewala, I am interested in BHMS admission at [College Name].\"",
        "The WhatsApp number should be configurable from the admin/settings panel rather than hard-coded throughout the frontend."
    ]),
    ("39. CONTACT PAGE", [
        "Include: Company information; Phone; Email; WhatsApp; Working hours; Contact form; Location/map where appropriate; Social links."
    ]),
    ("40. ABOUT PAGE", [
        "Explain: Padhaanewala; Mission; Vision; Services; College discovery; Student support; Counselling.",
        "Avoid legally risky or unverifiable claims such as guaranteed admission unless they are genuinely supportable."
    ]),
    ("41. LEGAL PAGES", [
        "Create: Privacy Policy; Terms & Conditions; Disclaimer; Cookie Policy where applicable; Refund/Cancellation Policy if payments are introduced."
    ]),
    ("42. API ARCHITECTURE", [
        "Use versioned APIs such as /api/v1/auth, /users, /colleges, /courses, /scholarships, /exams, /mock-tests, /reviews, /blog, /enquiries, /predictor, /ai and /admin.",
        "APIs must be documented, preferably through FastAPI OpenAPI/Swagger."
    ]),
    ("43. COLLEGE API", [
        "Example endpoints: GET /api/v1/colleges; GET /api/v1/colleges/{id}; POST /api/v1/colleges; PUT /api/v1/colleges/{id}; DELETE /api/v1/colleges/{id}.",
        "Public users should only have access to appropriate public endpoints. Admin endpoints must require authentication and authorization."
    ]),
    ("44. SEARCH API", [
        "Example: GET /api/v1/colleges.",
        "Parameters: course; state; district; city; college_type; min_fee; max_fee; rating; page; limit; sort.",
        "Do not return thousands of records in one response. Use pagination."
    ]),
    ("45. DATABASE ENTITIES", [
        "Minimum suggested entities: users; student_profiles; admins; counsellors; colleges; college_courses; courses; universities; locations; fees; admissions; scholarships; exams; exam_dates; mock_tests; questions; answers; test_attempts; reviews; blogs; categories; enquiries; lead_notes; notifications; saved_colleges; faqs; media; seo_metadata; audit_logs.",
        "Engineer may normalize this structure as required."
    ]),
    ("46. IMPORT SYSTEM", [
        "Admin should be able to upload college/course information using CSV/Excel.",
        "Example columns: College Name; State; District; City; University; Course; Fees; Website; Phone; Email; Address.",
        "Process: Upload → Validate → Show Errors → Preview → Confirm → Import."
    ]),
    ("47. DUPLICATE DETECTION", [
        "The import system should detect possible duplicate colleges.",
        "Do not create duplicate records simply because names differ slightly. Similar institutional names must be checked before creating separate records."
    ]),
    ("48. ANALYTICS", [
        "Website: Visitors; Searches; College views; Course views; Predictor usage.",
        "Students: Registrations; Active users.",
        "Leads: New leads; Contacted; Converted.",
        "Content: Most viewed colleges; Most searched courses; Popular scholarships; Popular articles."
    ]),
    ("49. LEAD SOURCE TRACKING", [
        "Store source for each enquiry.",
        "Examples: Homepage; College page; Course page; Blog; Scholarship; Predictor; WhatsApp; Advertisement; Organic search.",
        "This is important for measuring marketing performance."
    ]),
    ("50. UTM TRACKING", [
        "Support utm_source, utm_medium, utm_campaign and utm_content.",
        "Where practical, associate campaign information with the lead."
    ]),
    ("51. GOOGLE SERVICES", [
        "Configure Google Analytics, Google Search Console and Google Tag Manager if needed.",
        "Track searches, college views, predictor usage, enquiries, registrations, WhatsApp clicks, phone clicks and scholarship clicks."
    ]),
    ("52. SECURITY", [
        "Mandatory security requirements: HTTPS; secure authentication; password hashing; secure sessions/JWT; input validation; SQL injection protection; XSS protection; rate limiting; API authorization; file-upload validation; secure HTTP headers; admin 2FA where possible; audit logs.",
        "Never store passwords in plain text or expose API keys in frontend code."
    ]),
    ("53. ENVIRONMENT VARIABLES", [
        "Use secure environment variables.",
        "Examples: DATABASE_URL; JWT_SECRET; AI_API_KEY; EMAIL_API_KEY; STORAGE_ACCESS_KEY; STORAGE_SECRET.",
        "Create separate configurations for development, staging and production."
    ]),
    ("54. PERFORMANCE", [
        "Optimize with server-side rendering/static generation where appropriate, image optimization, lazy loading, code splitting, CDN, API caching, database indexing and Redis caching.",
        "Avoid unnecessary animations and oversized images."
    ]),
    ("55. MOBILE RESPONSIVENESS", [
        "The website must work properly on Android, iPhone, tablets, laptops and desktop.",
        "Mobile should not simply be a scaled-down desktop interface. Important actions must be easy to use with touch."
    ]),
    ("56. ACCESSIBILITY", [
        "Implement proper heading hierarchy, alt text, accessible buttons, form labels, keyboard navigation, good contrast and screen-reader-friendly structure."
    ]),
    ("57. ERROR HANDLING", [
        "Do not expose technical errors to users.",
        "Instead of showing a database or server error, show: \"Something went wrong. Please try again.\"",
        "Technical details should be stored in server logs."
    ]),
    ("58. LOGGING AND MONITORING", [
        "Maintain application logs for API errors, authentication failures, admin changes, AI failures, database errors and payment errors if introduced.",
        "Set up monitoring/alerts for critical production failures."
    ]),
    ("59. BACKUPS", [
        "Database backup must be automated.",
        "Recommended: Daily backups; retention policy; off-site backup; periodic restoration testing.",
        "A backup must be tested to ensure it can actually be restored."
    ]),
    ("60. GIT AND VERSION CONTROL", [
        "Use Git.",
        "Suggested structure: main; develop; feature/*; bugfix/*.",
        "Use pull requests/code review where practical.",
        "Never commit production secrets."
    ]),
    ("61. TESTING", [
        "Frontend: Test forms, navigation, search, mobile layouts, authentication and college pages.",
        "Backend: Test APIs, authentication, authorization, database operations, search, predictor and enquiries.",
        "End-to-end: Student → Search → College → Enquiry; Student → Predictor → Results → Enquiry; Student → Mock Test → Submit → Result; Student → Scholarship → Official application information; Admin → Add College → Publish → Public college page."
    ]),
    ("62. ADMIN CONTENT WORKFLOW", [
        "For important information use: Draft → Review → Publish.",
        "This is especially useful for Fees, Admission dates, Scholarships and Examination dates."
    ]),
    ("63. PHASED DEVELOPMENT", [
        "Phase 1 — Core Platform: Homepage; College database; College search; College pages; Course pages; Admin panel; CMS; Enquiry system; Authentication; Basic SEO.",
        "Phase 2 — Student Features: Student dashboard; Saved colleges; College comparison; Scholarship finder; Exam notifications; Reviews.",
        "Phase 3 — AI: AI College Predictor; AI Education Assistant; AI college comparison.",
        "Phase 4 — Mock Tests: Question bank; Test engine; Timer; Results; Performance analytics.",
        "Phase 5 — Business/CRM: Counsellor dashboard; Lead assignment; Follow-ups; CRM; Marketing attribution; Advanced analytics; Automated communication."
    ]),
    ("64. FINAL SYSTEM ARCHITECTURE", [
        "Recommended architecture: Students/Admin → Next.js Frontend/Admin Dashboard → REST APIs → FastAPI → PostgreSQL / Redis / AI / Storage / Search.",
        "Database contains Colleges, Courses, Scholarships, Exams, Students, Reviews, Mock Tests, Leads and Content.",
        "Architecture must be scalable and maintainable."
    ]),
    ("65. MOST IMPORTANT DEVELOPMENT RULES", [
        "1. Do not hard-code college/course/fee/scholarship data into React components.",
        "2. All major content must be database-driven.",
        "3. Admin must be able to update content without modifying source code.",
        "4. Frontend and backend must communicate through documented APIs.",
        "5. AI must not be allowed to freely invent factual college information.",
        "6. Security must be implemented from the beginning, not added at the end.",
        "7. Mobile responsiveness must be built from the beginning.",
        "8. SEO must be considered during architecture, not after development.",
        "9. All important data should have a source and/or verification date where applicable.",
        "10. The platform must be designed so additional colleges, courses, states, users and features can be added without redesigning the entire platform."
    ]),
    ("66. DEVELOPER DELIVERABLES", [
        "At completion, the engineer/team must provide: Complete source code; Frontend code; Backend code; Database schema; API documentation; Admin credentials; Deployment documentation; Environment-variable documentation; Database backup procedure; Git repository; Testing report; Production deployment; Staging deployment; Domain/DNS configuration documentation; Third-party service documentation.",
        "Padhaanewala should retain appropriate ownership/access to source code, hosting account, domain, database and relevant third-party accounts."
    ]),
    ("67. DEFINITION OF READY FOR LAUNCH", [
        "The website is considered ready for production only when: Homepage works; College search works; College pages work; Course pages work; Admin works; Content can be edited without code; Enquiries reach admin; Authentication works; Mobile version works; SEO basics are configured; HTTPS works; Database backup works; Security testing is completed; Major APIs are tested; Error handling works; Analytics is configured; Sitemap is available; Production deployment is stable."
    ]),
    ("68. FUTURE EXPANSION", [
        "The architecture should allow future features such as Online counselling; Paid counselling; Application tracking; Student document management; Online payments; Premium memberships; College advertising; Sponsored listings; College CRM; Partner college dashboard; Student application dashboard; Scholarship application tracking; AI career counselling; Personalized course recommendations; Mobile application; Push notifications; Multilingual support.",
        "These features do not necessarily need to be built in Version 1, but the architecture should not prevent them from being added later."
    ]),
    ("69. FINAL INSTRUCTION TO ENGINEER", [
        "Padhaanewala should be developed as a scalable education technology platform, not merely as a website.",
        "Priority: Reliable data + excellent search + strong college pages + admin control + lead generation + SEO + scalable backend.",
        "AI, mock tests and advanced features should be built on top of this foundation.",
        "The founder/admin must be able to manage the platform without depending on the developer for routine content changes.",
        "The final product should be fast, secure, mobile-friendly, SEO-friendly, scalable and easy to maintain."
    ])
]

for title, bullets in sections:
    doc.add_heading(title, level=2)
    for b in bullets:
        add_bullet(b)

doc.add_page_break()

# ============================================
# SECTION B: RAG ADDITION
# ============================================
doc.add_heading('SECTION B: RAG (RETRIEVAL-AUGMENTED GENERATION)', level=1)

doc.add_heading('B.1. What is RAG?', level=2)
add_body("RAG (Retrieval-Augmented Generation) is an AI architecture pattern that combines information retrieval with language model generation. Instead of relying solely on the AI model's training data, RAG retrieves relevant documents from a database before generating a response.")

doc.add_heading('B.2. Why RAG for Padhaanewala?', level=2)
add_bullet("AI College Predictor needs accurate, up-to-date college data")
add_bullet("AI Education Assistant must provide verified information, not hallucinated facts")
add_bullet("Natural Language Search requires understanding student queries and mapping to structured data")
add_bullet("College Comparison needs real database information for accurate comparisons")
add_bullet("Prevents AI from making up college names, fees, or admission criteria")

doc.add_heading('B.3. RAG Architecture Flow', level=2)
add_body("Complete RAG Pipeline:")
add_bullet("Step 1: Student Query → \"BHMS colleges in Karnataka under 5 lakh with hostel\"")
add_bullet("Step 2: Query Preprocessing → Extract: Course=BHMS, State=Karnataka, MaxFee=500000, Hostel=Yes")
add_bullet("Step 3: Structured DB Search → PostgreSQL filter with extracted parameters")
add_bullet("Step 4: Embedding Generation → Convert query to vector using OpenAI/Cohere embeddings")
add_bullet("Step 5: Vector Search → pgvector similarity search in document_embeddings table")
add_bullet("Step 6: Context Assembly → Combine DB results + vector results into LLM context")
add_bullet("Step 7: LLM Processing → Send context + query to GPT-4/Claude for response generation")
add_bullet("Step 8: Response + Sources → Answer with source citations and verification disclaimer")

doc.add_heading('B.4. Database Changes for RAG', level=2)
add_body("New Table: document_embeddings")
add_bullet("id: UUID (Primary Key)")
add_bullet("entity_type: VARCHAR(50) — college, course, scholarship, faq")
add_bullet("entity_id: UUID — reference to the source entity")
add_bullet("chunk_text: TEXT — the text chunk being embedded")
add_bullet("embedding: VECTOR(1536) — pgvector extension for vector storage")
add_bullet("metadata: JSONB — additional context information")
add_bullet("created_at: TIMESTAMP")
add_bullet("updated_at: TIMESTAMP")

add_body("Index for fast vector search:")
add_bullet("CREATE INDEX ON document_embeddings USING ivfflat (embedding vector_cosine_ops) WITH (lists = 100);")

doc.add_heading('B.5. AI College Predictor with RAG', level=2)
add_body("Input Processing:")
add_bullet("Course, Entrance Exam, Rank/Score, Category, State, City, Budget, College Type, Hostel")
add_body("RAG Flow:")
add_bullet("1. Extract structured filters from student inputs")
add_bullet("2. Query PostgreSQL for eligible colleges matching criteria")
add_bullet("3. Generate embedding for student's preference description")
add_bullet("4. Vector search for colleges with similar student satisfaction patterns")
add_bullet("5. Combine results and rank using AI")
add_bullet("6. Categorize: Highly Suitable / Possible / Reach")
add_bullet("7. Generate explanation based on verified college data")
add_bullet("8. Add disclaimer: \"Results are estimates, not guaranteed admissions\"")

doc.add_heading('B.6. AI Education Assistant with RAG', level=2)
add_body("Question Processing:")
add_bullet("Preprocess question to extract intent and entities")
add_bullet("Example: \"What is BHMS?\" → Intent: Course Information, Entity: BHMS")
add_body("RAG Flow:")
add_bullet("1. Search PostgreSQL for BHMS course data")
add_bullet("2. Vector search for similar educational content in embeddings")
add_bullet("3. Retrieve FAQs related to BHMS")
add_bullet("4. Assemble context with all retrieved information")
add_bullet("5. Generate comprehensive answer using LLM")
add_bullet("6. Cite sources: \"Based on data from College X, Course Y\"")
add_bullet("7. Add disclaimer: \"Information may change, please verify with institution\"")

doc.add_heading('B.7. Natural Language Search with RAG', level=2)
add_body("Query Parsing:")
add_bullet("\"Nursing colleges near Bangalore under 5 lakh\"")
add_bullet("→ Course: B.Sc Nursing")
add_bullet("→ Location: Bangalore (within 100km radius)")
add_bullet("→ Max Fee: 500000")
add_body("RAG Flow:")
add_bullet("1. Parse natural language query into structured filters")
add_bullet("2. Execute PostgreSQL search with filters")
add_bullet("3. If results are insufficient, expand search using vector similarity")
add_bullet("4. Rank results by relevance, rating, and proximity")
add_bullet("5. Return structured results with explanations")

doc.add_heading('B.8. Embedding Generation Strategy', level=2)
add_body("Data Chunking:")
add_bullet("College: Name, Location, Courses, Fees, Facilities, Reviews (separate chunks)")
add_bullet("Course: Name, Duration, Eligibility, Career Info (separate chunks)")
add_bullet("Scholarship: Name, Eligibility, Amount, Deadline (separate chunks)")
add_bullet("FAQ: Question + Answer pairs (individual chunks)")
add_body("Embedding Model:")
add_bullet("Primary: text-embedding-3-small (OpenAI) — 1536 dimensions")
add_bullet("Alternative: embed-english-v3.0 (Cohere) — 1024 dimensions")
add_bullet("Local: all-MiniLM-L6-v2 (for cost optimization)")

doc.add_heading('B.9. Vector Search Configuration', level=2)
add_bullet("Index Type: IVFFlat (Inverted File Index) — good balance of speed and accuracy")
add_bullet("Similarity Metric: Cosine Similarity — best for semantic search")
add_bullet("Lists: 100 (for <100k documents) or 1000 (for >100k documents)")
add_bullet("Top-K Retrieval: 5-10 most similar chunks")

doc.add_heading('B.10. AI Guardrails', level=2)
add_bullet("No Hallucination: AI must only use retrieved database information")
add_bullet("Source Citation: Every AI response must cite its data sources")
add_bullet("Uncertainty Disclosure: If information is uncertain, explicitly state it")
add_bullet("Verification Reminder: Always remind users to verify with institutions")
add_bullet("No Guarantees: Never promise admissions or guaranteed outcomes")
add_bullet("Data Freshness: Indicate when data was last verified")

doc.add_heading('B.11. RAG Implementation Checklist', level=2)
add_bullet("☐ Install pgvector extension in PostgreSQL")
add_bullet("☐ Create document_embeddings table")
add_bullet("☐ Set up embedding generation service (OpenAI/Cohere)")
add_bullet("☐ Create Celery task for embedding generation on data changes")
add_bullet("☐ Implement vector search API endpoint")
add_bullet("☐ Build RAG pipeline in FastAPI")
add_bullet("☐ Integrate with AI College Predictor")
add_bullet("☐ Integrate with AI Education Assistant")
add_bullet("☐ Integrate with Natural Language Search")
add_bullet("☐ Add source citation to all AI responses")
add_bullet("☐ Add disclaimers to all AI outputs")
add_bullet("☐ Test with sample queries and validate accuracy")
add_bullet("☐ Monitor AI response quality and hallucination rate")

doc.add_page_break()

# ============================================
# SECTION C: ARCHITECTURE IMPROVEMENTS
# ============================================
doc.add_heading('SECTION C: ARCHITECTURE IMPROVEMENTS', level=1)

doc.add_heading('C.1. Infrastructure Architecture', level=2)
add_body("Updated Production Architecture:")
add_bullet("Internet → Cloudflare CDN (DDoS Protection + WAF)")
add_bullet("Cloudflare → Load Balancer (Nginx)")
add_bullet("Load Balancer → Next.js (3 instances)")
add_bullet("Next.js → Load Balancer")
add_bullet("Load Balancer → FastAPI (3 instances)")
add_bullet("FastAPI → PostgreSQL (Primary + Replica)")
add_bullet("FastAPI → Redis (Cluster)")
add_bullet("FastAPI → Object Storage (S3/R2)")
add_bullet("FastAPI → Search (PostgreSQL → OpenSearch later)")

doc.add_heading('C.2. Background Task Queue (Celery)', level=2)
add_body("Why Needed:")
add_bullet("CSV import 1000 colleges = 5 minutes → Frontend stuck without async")
add_bullet("AI processing = 10-30 seconds → User waiting")
add_bullet("Email sending = 2-5 seconds × 1000 users = slow")
add_bullet("Image optimization = heavy process")

add_body("Celery Configuration:")
add_bullet("Broker: Redis")
add_bullet("Backend: Redis")
add_bullet("Worker: 4 concurrent workers (production)")

add_body("Task Types:")
add_bullet("import_tasks.py — CSV/Excel bulk import jobs")
add_bullet("ai_tasks.py — AI processing (predictor, assistant, comparison)")
add_bullet("email_tasks.py — Email sending (OTP, notifications, marketing)")
add_bullet("sms_tasks.py — SMS sending (OTP, confirmations)")
add_bullet("image_tasks.py — Image optimization (resize, compress, WebP)")
add_bullet("sitemap_tasks.py — Sitemap generation")
add_bullet("analytics_tasks.py — Analytics aggregation")
add_bullet("backup_tasks.py — Database backup scheduling")

doc.add_heading('C.3. Email Service Integration', level=2)
add_body("Provider: SendGrid (recommended) or AWS SES")
add_body("Email Templates:")
add_bullet("welcome — Registration confirmation")
add_bullet("otp_verification — Mobile/Email OTP")
add_bullet("password_reset — Password reset link")
add_bullet("enquiry_confirmation — Enquiry submitted successfully")
add_bullet("lead_assigned — Counsellor assigned to lead")
add_bullet("lead_followup — Follow-up reminder")
add_bullet("weekly_digest — Weekly education news")

doc.add_heading('C.4. SMS Service Integration', level=2)
add_body("Provider: MSG91 (best for India) or Twilio")
add_body("SMS Templates:")
add_bullet("otp — 6-digit OTP for verification")
add_bullet("enquiry_received — Enquiry confirmation")
add_bullet("counsellor_call — Counsellor will call soon")

doc.add_heading('C.5. WhatsApp Business API', level=2)
add_body("Provider: Twilio / WATI / AiSensy")
add_body("Contextual Messages:")
add_bullet("college_inquiry: \"Hello Padhaanewala, I am interested in {course} admission at {college_name}.\"")
add_bullet("general: \"Hello Padhaanewala, I need guidance about education.\"")
add_bullet("scholarship: \"Hello Padhaanewala, I want to know about {scholarship_name}.\"")

doc.add_heading('C.6. Google Maps Integration', level=2)
add_body("APIs Needed:")
add_bullet("Geocoding API — College address → Latitude/Longitude")
add_bullet("Maps Embed API — Display college location on map")
add_bullet("Places API — Auto-suggest locations")

add_body("Database Fields to Add:")
add_bullet("latitude: DECIMAL(10, 8)")
add_bullet("longitude: DECIMAL(11, 8)")
add_bullet("google_maps_url: TEXT")
add_bullet("google_place_id: VARCHAR(255)")

doc.add_heading('C.7. Image Optimization Pipeline', level=2)
add_body("Flow: Admin Upload → Validate → Optimize → Upload to S3 → Store URL → Serve via CDN")
add_body("Optimization Steps:")
add_bullet("Resize to multiple sizes: Thumbnail (150x150), Card (400x300), Detail (800x600)")
add_bullet("Compress to 80% quality")
add_bullet("Convert to WebP format (modern, smaller)")
add_bullet("Generate blur placeholder for lazy loading")
add_body("File Validation:")
add_bullet("Allowed image types: .jpg, .jpeg, .png, .webp")
add_bullet("Max file size: 5MB for images")
add_bullet("MIME type verification")
add_bullet("Malware scanning for documents")

doc.add_heading('C.8. Monitoring & Logging', level=2)
add_body("Components:")
add_bullet("Structured Logging — JSON format for all application logs")
add_bullet("Sentry — Real-time error tracking with stack traces")
add_bullet("Performance Monitoring — Datadog / New Relic for API response times")
add_bullet("Uptime Monitoring — UptimeRobot / Pingdom for 24/7 checks")
add_bullet("Alerting — Slack / Email / PagerDuty for critical issues")

add_body("Log Categories:")
add_bullet("api — API request/response logs")
add_bullet("auth — Authentication attempts and failures")
add_bullet("ai — AI processing logs and failures")
add_bullet("db — Database query logs")
add_bullet("admin — Admin action audit logs")
add_bullet("security — Security event logs")

doc.add_heading('C.9. Backup & Recovery', level=2)
add_body("PostgreSQL Backup:")
add_bullet("Daily Full Backup → S3 (off-site) at 00:00")
add_bullet("WAL Archive → Every 6 hours → S3 (point-in-time recovery)")
add_bullet("Weekly Restoration Test → Restore to staging, verify integrity")
add_bullet("Retention: Daily (7 days), Weekly (4 weeks), Monthly (12 months)")

add_body("Redis Backup:")
add_bullet("RDB Snapshot → Every 15 minutes → S3")
add_bullet("AOF Backup → Every hour → S3")

add_body("Object Storage:")
add_bullet("S3/R2 Versioning → Enabled")
add_bullet("Cross-Region Replication → Secondary bucket")

doc.add_heading('C.10. Security Enhancements', level=2)
add_body("Request Flow with Security:")
add_bullet("1. Client Request → Cloudflare (DDoS Protection + WAF)")
add_bullet("2. Rate Limiter (Redis) — IP: 100 req/min, User: 1000 req/hour")
add_bullet("3. Security Headers (Helmet.js) — X-Content-Type, X-Frame, HSTS, CSP")
add_bullet("4. CORS Validation — Only padhaanewala.in allowed")
add_bullet("5. Input Validation (Pydantic) — SQL injection, XSS, length, type checks")
add_bullet("6. Authentication (JWT) — Access: 15 min, Refresh: 7 days, Admin 2FA")
add_bullet("7. Authorization (RBAC) — Role-based access control")
add_bullet("8. Audit Logging — Who, What, When, Where, Result")

doc.add_heading('C.11. SEO Enhancements', level=2)
add_body("Sitemap Generator:")
add_bullet("Auto-generate XML sitemap with all pages")
add_bullet("Include: Static pages, College pages, Course pages, Blog pages")
add_bullet("Priority: Homepage (1.0), Colleges (0.9), Courses (0.9), Blog (0.7)")
add_bullet("Update frequency: Daily for dynamic content")

add_body("Schema.org Markup:")
add_bullet("College: EducationalOrganization schema")
add_bullet("Course: Course schema")
add_bullet("Scholarship: Scholarship schema")
add_bullet("Blog: Article schema")
add_bullet("FAQ: FAQPage schema")
add_bullet("BreadcrumbList: Navigation schema")

doc.add_heading('C.12. Complete Updated Architecture Diagram', level=2)
add_body("Production Architecture:")
add_bullet("Internet → Cloudflare CDN → Load Balancer → Next.js (x3)")
add_bullet("Next.js → Load Balancer → FastAPI (x3)")
add_bullet("FastAPI → PostgreSQL (Primary + Replica)")
add_bullet("FastAPI → Redis (Cluster)")
add_bullet("FastAPI → Object Storage (S3/R2)")
add_bullet("FastAPI → Celery Worker (Background Tasks)")
add_bullet("Celery → Redis (Broker)")
add_bullet("Celery → Email/SMS/WhatsApp Services")
add_bullet("Celery → Image Optimization")
add_bullet("Celery → Embedding Generation (RAG)")
add_bullet("Monitoring → Sentry + Datadog + UptimeRobot")
add_bullet("Backup → S3 (Daily + WAL)")
add_bullet("CI/CD → GitHub Actions → Docker → Production")

doc.add_page_break()

# ============================================
# SECTION D: DATABASE DESIGN
# ============================================
doc.add_heading('SECTION D: DATABASE DESIGN', level=1)

doc.add_heading('D.1. Core Tables', level=2)

tables = [
    ("users", "id, email, phone, password_hash, role, is_active, created_at, updated_at"),
    ("student_profiles", "id, user_id, name, education, course_interest, preferred_state, preferred_city, budget, created_at"),
    ("admins", "id, user_id, role, permissions, created_at"),
    ("counsellors", "id, user_id, name, phone, email, max_leads, created_at"),
    ("colleges", "id, name, slug, official_name, type, is_government, university_id, state, district, city, address, pincode, website, email, phone, established_year, accreditation, recognition, description, logo_url, cover_url, latitude, longitude, data_source, verification_status, last_verified_date, verified_by, status, created_at, updated_at"),
    ("courses", "id, name, slug, degree, duration, description, eligibility, entrance_exam, admission_procedure, career_info, seo_title, seo_description, status, created_at, updated_at"),
    ("college_courses", "id, college_id, course_id, fees_id, admission_info, seats, created_at"),
    ("universities", "id, name, slug, state, website, description"),
    ("locations", "id, state, district, city, latitude, longitude"),
    ("fees", "id, college_course_id, tuition_fee, hostel_fee, exam_fee, other_charges, total_approximate, fee_period, is_approximate, created_at"),
    ("admissions", "id, college_course_id, eligibility, entrance_exam, cutoff, admission_process, important_dates, created_at"),
    ("scholarships", "id, name, slug, provider, is_government, eligibility, state, course, income_criteria, amount, deadline, documents, application_procedure, official_link, status, last_verified_date, created_at"),
    ("exams", "id, name, slug, conducting_authority, eligibility, application_start, application_deadline, exam_date, admit_card_date, result_date, official_website, notification_url, status, created_at"),
    ("mock_tests", "id, name, exam_id, subject, difficulty, total_questions, time_limit, instructions, status, created_at"),
    ("questions", "id, mock_test_id, question_text, question_type, options, correct_answer, explanation, marks, created_at"),
    ("test_attempts", "id, user_id, mock_test_id, answers, score, percentage, correct, incorrect, unattempted, time_taken, topic_wise_performance, rank, percentile, started_at, submitted_at"),
    ("reviews", "id, user_id, college_id, course, year, rating, review_text, images, status, moderation_notes, created_at"),
    ("blogs", "id, title, slug, content, featured_image, category_id, author_id, meta_title, meta_description, canonical_url, status, published_at, created_at"),
    ("categories", "id, name, slug, description, parent_id, type"),
    ("enquiries", "id, user_id, name, phone, email, course, preferred_college, state, city, qualification, message, source, utm_source, utm_medium, utm_campaign, utm_content, status, assigned_counsellor_id, created_at"),
    ("lead_notes", "id, enquiry_id, counsellor_id, note, created_at"),
    ("notifications", "id, user_id, title, message, type, is_read, created_at"),
    ("saved_colleges", "id, user_id, college_id, created_at"),
    ("faqs", "id, entity_type, entity_id, question, answer, order, status, created_at"),
    ("media", "id, filename, original_name, mime_type, size, url, alt_text, entity_type, entity_id, uploaded_by, created_at"),
    ("seo_metadata", "id, entity_type, entity_id, title, description, canonical_url, og_title, og_description, og_image, schema_markup, created_at"),
    ("audit_logs", "id, user_id, action, entity_type, entity_id, old_value, new_value, ip_address, user_agent, created_at"),
    ("document_embeddings", "id, entity_type, entity_id, chunk_text, embedding, metadata, created_at, updated_at"),
]

add_table(["Table Name", "Key Columns"], tables)

doc.add_heading('D.2. Relationships', level=2)
add_bullet("users → student_profiles (1:1)")
add_bullet("users → admins (1:1)")
add_bullet("users → counsellors (1:1)")
add_bullet("users → reviews (1:many)")
add_bullet("users → saved_colleges (1:many)")
add_bullet("users → test_attempts (1:many)")
add_bullet("colleges → college_courses (1:many)")
add_bullet("courses → college_courses (1:many)")
add_bullet("college_courses → fees (1:1)")
add_bullet("college_courses → admissions (1:1)")
add_bullet("colleges → reviews (1:many)")
add_bullet("colleges → saved_colleges (1:many)")
add_bullet("colleges → faqs (1:many)")
add_bullet("colleges → media (1:many)")
add_bullet("enquiries → lead_notes (1:many)")
add_bullet("counsellors → enquiries (1:many)")
add_bullet("mock_tests → questions (1:many)")
add_bullet("mock_tests → test_attempts (1:many)")

doc.add_heading('D.3. Indexes', level=2)
add_bullet("colleges: slug, state, city, type, status")
add_bullet("courses: slug, status")
add_bullet("college_courses: college_id, course_id")
add_bullet("scholarships: state, course, deadline, status")
add_bullet("exams: exam_date, status")
add_bullet("reviews: college_id, status")
add_bullet("blogs: slug, category_id, status")
add_bullet("enquiries: status, assigned_counsellor_id")
add_bullet("document_embeddings: entity_type, entity_id")
add_bullet("document_embeddings: ivfflat index on embedding (vector)")

doc.add_page_break()

# ============================================
# SECTION E: API DESIGN
# ============================================
doc.add_heading('SECTION E: API DESIGN', level=1)

doc.add_heading('E.1. Authentication APIs', level=2)
apis = [
    ("POST", "/api/v1/auth/register", "Register new user", "No"),
    ("POST", "/api/v1/auth/login", "Login user", "No"),
    ("POST", "/api/v1/auth/otp/send", "Send OTP", "No"),
    ("POST", "/api/v1/auth/otp/verify", "Verify OTP", "No"),
    ("POST", "/api/v1/auth/refresh", "Refresh token", "Yes"),
    ("POST", "/api/v1/auth/logout", "Logout user", "Yes"),
    ("POST", "/api/v1/auth/password/forgot", "Forgot password", "No"),
    ("POST", "/api/v1/auth/password/reset", "Reset password", "Yes"),
]
add_table(["Method", "Endpoint", "Description", "Auth Required"], apis)

doc.add_heading('E.2. College APIs', level=2)
apis = [
    ("GET", "/api/v1/colleges", "List colleges (with filters)", "No"),
    ("GET", "/api/v1/colleges/{id}", "Get college details", "No"),
    ("GET", "/api/v1/colleges/{id}/courses", "Get college courses", "No"),
    ("GET", "/api/v1/colleges/{id}/reviews", "Get college reviews", "No"),
    ("GET", "/api/v1/colleges/{id}/gallery", "Get college gallery", "No"),
    ("GET", "/api/v1/colleges/{id}/faqs", "Get college FAQs", "No"),
    ("POST", "/api/v1/colleges", "Create college", "Yes (Admin)"),
    ("PUT", "/api/v1/colleges/{id}", "Update college", "Yes (Admin)"),
    ("DELETE", "/api/v1/colleges/{id}", "Delete college", "Yes (Super Admin)"),
    ("POST", "/api/v1/colleges/import", "Import colleges (CSV)", "Yes (Admin)"),
]
add_table(["Method", "Endpoint", "Description", "Auth Required"], apis)

doc.add_heading('E.3. Course APIs', level=2)
apis = [
    ("GET", "/api/v1/courses", "List courses", "No"),
    ("GET", "/api/v1/courses/{id}", "Get course details", "No"),
    ("GET", "/api/v1/courses/{id}/colleges", "Get colleges offering course", "No"),
    ("POST", "/api/v1/courses", "Create course", "Yes (Admin)"),
    ("PUT", "/api/v1/courses/{id}", "Update course", "Yes (Admin)"),
    ("DELETE", "/api/v1/courses/{id}", "Delete course", "Yes (Super Admin)"),
]
add_table(["Method", "Endpoint", "Description", "Auth Required"], apis)

doc.add_heading('E.4. Search APIs', level=2)
apis = [
    ("GET", "/api/v1/search", "Universal search", "No"),
    ("GET", "/api/v1/search/colleges", "College search with filters", "No"),
    ("GET", "/api/v1/search/natural", "Natural language search", "No"),
    ("GET", "/api/v1/search/suggestions", "Search suggestions", "No"),
]
add_table(["Method", "Endpoint", "Description", "Auth Required"], apis)

doc.add_heading('E.5. AI APIs', level=2)
apis = [
    ("POST", "/api/v1/predictor", "AI College Predictor", "No"),
    ("POST", "/api/v1/ai/assistant", "AI Education Assistant", "No"),
    ("POST", "/api/v1/ai/compare", "AI College Comparison", "No"),
    ("POST", "/api/v1/ai/recommend", "AI Recommendations", "No"),
]
add_table(["Method", "Endpoint", "Description", "Auth Required"], apis)

doc.add_heading('E.6. Student APIs', level=2)
apis = [
    ("GET", "/api/v1/users/me", "Get current user profile", "Yes"),
    ("PUT", "/api/v1/users/me", "Update profile", "Yes"),
    ("GET", "/api/v1/users/me/saved-colleges", "Get saved colleges", "Yes"),
    ("POST", "/api/v1/users/me/saved-colleges", "Save college", "Yes"),
    ("DELETE", "/api/v1/users/me/saved-colleges/{id}", "Remove saved college", "Yes"),
    ("GET", "/api/v1/users/me/test-history", "Get test history", "Yes"),
    ("GET", "/api/v1/users/me/enquiries", "Get enquiries", "Yes"),
]
add_table(["Method", "Endpoint", "Description", "Auth Required"], apis)

doc.add_heading('E.7. Mock Test APIs', level=2)
apis = [
    ("GET", "/api/v1/mock-tests", "List mock tests", "No"),
    ("GET", "/api/v1/mock-tests/{id}", "Get mock test details", "No"),
    ("POST", "/api/v1/mock-tests/{id}/start", "Start test attempt", "Yes"),
    ("POST", "/api/v1/mock-tests/{id}/submit", "Submit test", "Yes"),
    ("GET", "/api/v1/mock-tests/{id}/result", "Get test result", "Yes"),
    ("GET", "/api/v1/mock-tests/{id}/solutions", "Get solutions", "Yes"),
]
add_table(["Method", "Endpoint", "Description", "Auth Required"], apis)

doc.add_heading('E.8. Enquiry/Lead APIs', level=2)
apis = [
    ("POST", "/api/v1/enquiries", "Submit enquiry", "No"),
    ("GET", "/api/v1/enquiries", "List enquiries", "Yes (Admin)"),
    ("GET", "/api/v1/enquiries/{id}", "Get enquiry details", "Yes (Admin)"),
    ("PUT", "/api/v1/enquiries/{id}", "Update enquiry", "Yes (Admin)"),
    ("POST", "/api/v1/enquiries/{id}/notes", "Add note to enquiry", "Yes (Counsellor)"),
    ("PUT", "/api/v1/enquiries/{id}/assign", "Assign counsellor", "Yes (Admin)"),
]
add_table(["Method", "Endpoint", "Description", "Auth Required"], apis)

doc.add_heading('E.9. Admin APIs', level=2)
apis = [
    ("GET", "/api/v1/admin/dashboard", "Get dashboard data", "Yes (Admin)"),
    ("GET", "/api/v1/admin/analytics", "Get analytics", "Yes (Admin)"),
    ("GET", "/api/v1/admin/audit-logs", "Get audit logs", "Yes (Super Admin)"),
    ("POST", "/api/v1/admin/banners", "Create banner", "Yes (Admin)"),
    ("PUT", "/api/v1/admin/banners/{id}", "Update banner", "Yes (Admin)"),
    ("GET", "/api/v1/admin/leads", "Get all leads", "Yes (Admin)"),
    ("GET", "/api/v1/admin/counsellors", "List counsellors", "Yes (Admin)"),
    ("POST", "/api/v1/admin/counsellors", "Create counsellor", "Yes (Super Admin)"),
]
add_table(["Method", "Endpoint", "Description", "Auth Required"], apis)

doc.add_page_break()

# ============================================
# SECTION F: IMPLEMENTATION PLAN
# ============================================
doc.add_heading('SECTION F: IMPLEMENTATION PLAN (9 MONTHS)', level=1)

months = [
    ("Month 1: Foundation", [
        "Week 1-2: Project Setup",
        "  - Next.js + FastAPI + PostgreSQL initialization",
        "  - Docker setup for all services",
        "  - CI/CD pipeline (GitHub Actions)",
        "  - Authentication system (JWT + OTP)",
        "  - Environment variables configuration",
        "Week 3-4: Database Design",
        "  - All 45+ tables created",
        "  - Migrations with Alembic",
        "  - Seed data for development",
        "  - pgvector extension installed",
        "  - Basic API endpoints"
    ]),
    ("Month 2: Core Features", [
        "Week 5-6: College System",
        "  - College CRUD (Admin API)",
        "  - College pages (Frontend)",
        "  - College search with filters",
        "  - CSV import functionality",
        "  - College gallery upload",
        "Week 7-8: Course System",
        "  - Course CRUD",
        "  - Course pages",
        "  - College-Course relationships",
        "  - Fee structure management"
    ]),
    ("Month 3: Content & Search", [
        "Week 9-10: Scholarships + Exams",
        "  - Scholarship database + pages",
        "  - Exam database + pages",
        "  - Deadline tracking",
        "  - Notification system",
        "Week 11-12: Search Enhancement",
        "  - Natural language search (basic)",
        "  - Advanced filter UI",
        "  - Search results pagination",
        "  - Search suggestions/autocomplete"
    ]),
    ("Month 4: Student Features", [
        "Week 13-14: Student System",
        "  - Registration/Login (OTP + Email)",
        "  - Student dashboard",
        "  - Saved colleges functionality",
        "  - Profile management",
        "Week 15-16: Reviews + Comparison",
        "  - Review submission + moderation",
        "  - College comparison tool",
        "  - College gallery enhancement",
        "  - FAQ system"
    ]),
    ("Month 5: AI Features", [
        "Week 17-18: RAG Setup",
        "  - pgvector configuration",
        "  - Embedding generation service",
        "  - document_embeddings table",
        "  - Vector search API",
        "  - Celery task for embedding on data changes",
        "Week 19-20: AI Features",
        "  - AI College Predictor with RAG",
        "  - AI Education Assistant with RAG",
        "  - AI College Comparison",
        "  - Source citation in AI responses"
    ]),
    ("Month 6: Admin & CMS", [
        "Week 21-22: Admin Panel",
        "  - All 18 admin modules",
        "  - Role-based access control",
        "  - Media management",
        "  - Bulk operations",
        "Week 23-24: CMS",
        "  - Blog system",
        "  - FAQ management",
        "  - Banner management",
        "  - SEO metadata management"
    ]),
    ("Month 7: Mock Tests", [
        "Week 25-26: Question Bank",
        "  - Question CRUD (Admin)",
        "  - Question import (CSV)",
        "  - Question categories/difficulty",
        "  - Question validation",
        "Week 27-28: Test Engine",
        "  - Test creation (Admin)",
        "  - Test taking interface (Student)",
        "  - Timer + navigation",
        "  - Submission + scoring",
        "  - Results + analytics"
    ]),
    ("Month 8: CRM & Leads", [
        "Week 29-30: Lead System",
        "  - Enquiry forms (all pages)",
        "  - Lead management (Admin)",
        "  - Lead source tracking",
        "  - UTM tracking",
        "Week 31-32: Counsellor System",
        "  - Counsellor dashboard",
        "  - Lead assignment",
        "  - Follow-up system",
        "  - Notes system",
        "  - Contact status tracking"
    ]),
    ("Month 9: Testing & Launch", [
        "Week 33-34: Testing",
        "  - Unit tests (Backend + Frontend)",
        "  - Integration tests",
        "  - E2E tests (Playwright)",
        "  - Security testing",
        "  - Performance testing",
        "  - Mobile testing",
        "Week 35-36: Deployment",
        "  - Production server setup",
        "  - Domain/DNS configuration",
        "  - SSL certificate",
        "  - CDN setup (Cloudflare)",
        "  - Monitoring setup (Sentry)",
        "  - Backup setup",
        "  - Launch checklist verification",
        "  - 🚀 PRODUCTION LAUNCH"
    ])
]

for title, items in months:
    doc.add_heading(title, level=2)
    for item in items:
        if item.startswith("  -"):
            p = doc.add_paragraph(item[2:], style='List Bullet 2')
        else:
            add_bullet(item)

doc.add_heading('Team Requirements', level=2)
team = [
    ("Backend Developer (Python/FastAPI)", "1-2", "APIs, Database, Auth, AI integration"),
    ("Frontend Developer (Next.js/React)", "1-2", "UI, Pages, Components, Responsive"),
    ("Database Designer", "1", "Schema, Migrations, Optimization"),
    ("AI/ML Engineer (RAG)", "1", "Embeddings, Vector search, AI pipeline"),
    ("DevOps Engineer", "1", "Docker, CI/CD, Deployment, Monitoring"),
    ("UI/UX Designer", "1", "Design system, Wireframes, Prototypes"),
    ("Content/Data Entry Team", "2-3", "College data, Course data, Scholarships"),
]
add_table(["Role", "Count", "Responsibilities"], team)

doc.add_page_break()

# ============================================
# SECTION G: TESTING PLAN
# ============================================
doc.add_heading('SECTION G: TESTING PLAN', level=1)

doc.add_heading('G.1. Unit Tests', level=2)
add_bullet("Backend: All API endpoints, services, utilities")
add_bullet("Frontend: Components, hooks, utilities")
add_bullet("Database: Models, queries, migrations")

doc.add_heading('G.2. Integration Tests', level=2)
add_bullet("API + Database operations")
add_bullet("Authentication flow")
add_bullet("Authorization checks")
add_bullet("Search functionality")
add_bullet("AI pipeline")

doc.add_heading('G.3. E2E Tests (Playwright)', level=2)
add_bullet("Student → Search → College → Enquiry")
add_bullet("Student → Predictor → Results → Enquiry")
add_bullet("Student → Mock Test → Submit → Result")
add_bullet("Student → Scholarship → Official application info")
add_bullet("Admin → Add College → Publish → Public college page")
add_bullet("Admin → Manage Leads → Assign Counsellor")

doc.add_heading('G.4. Security Tests', level=2)
add_bullet("SQL injection attempts")
add_bullet("XSS attack attempts")
add_bullet("CSRF protection")
add_bullet("Rate limiting verification")
add_bullet("Authentication bypass attempts")
add_bullet("Authorization escalation attempts")
add_bullet("File upload validation")
add_bullet("Input validation")

doc.add_heading('G.5. Performance Tests', level=2)
add_bullet("Load testing: 1000 concurrent users")
add_bullet("Stress testing: 5000 concurrent users")
add_bullet("API response time: <200ms for most endpoints")
add_bullet("Database query performance")
add_bullet("Image load time: <2 seconds")

doc.add_heading('G.6. Mobile Testing', level=2)
add_bullet("Android (Chrome, Samsung Browser)")
add_bullet("iOS (Safari)")
add_bullet("Tablets (iPad, Android tablets)")
add_bullet("Touch interactions")
add_bullet("Responsive breakpoints: 320px, 768px, 1024px, 1440px")

doc.add_page_break()

# ============================================
# SECTION H: DEPLOYMENT GUIDE
# ============================================
doc.add_heading('SECTION H: DEPLOYMENT GUIDE', level=1)

doc.add_heading('H.1. Production Server Setup', level=2)
add_bullet("Cloud Provider: AWS / DigitalOcean / Vultr")
add_bullet("Minimum: 4 vCPU, 8GB RAM, 100GB SSD")
add_bullet("Recommended: 8 vCPU, 16GB RAM, 200GB SSD")
add_bullet("OS: Ubuntu 22.04 LTS")
add_bullet("Docker + Docker Compose installed")

doc.add_heading('H.2. Environment Variables', level=2)
envs = [
    ("DATABASE_URL", "postgresql://user:pass@host:5432/padhaanewala"),
    ("REDIS_URL", "redis://host:6379/0"),
    ("JWT_SECRET", "your-super-secret-jwt-key"),
    ("AI_API_KEY", "sk-... (OpenAI/Cohere)"),
    ("EMAIL_API_KEY", "SG... (SendGrid)"),
    ("SMS_API_KEY", "... (MSG91)"),
    ("WHATSAPP_API_KEY", "... (Twilio)"),
    ("STORAGE_ACCESS_KEY", "... (AWS S3/R2)"),
    ("STORAGE_SECRET", "... (AWS S3/R2)"),
    ("STORAGE_BUCKET", "padhaanewala-media"),
    ("SENTRY_DSN", "https://...@sentry.io/..."),
    ("NEXT_PUBLIC_API_URL", "https://api.padhaanewala.in"),
    ("NEXT_PUBLIC_GOOGLE_MAPS_KEY", "AIza..."),
]
add_table(["Variable", "Value"], envs)

doc.add_heading('H.3. DNS Configuration', level=2)
add_bullet("A Record: padhaanewala.in → Production Server IP")
add_bullet("CNAME: www.padhaanewala.in → padhaanewala.in")
add_bullet("CNAME: api.padhaanewala.in → Production Server IP")
add_bullet("MX Record: For email (if using custom domain)")
add_bullet("TXT Record: SPF, DKIM (for email authentication)")

doc.add_heading('H.4. SSL Certificate', level=2)
add_bullet("Provider: Cloudflare (Free) or Let's Encrypt")
add_bullet("Wildcard: *.padhaanewala.in")
add_bullet("Auto-renewal configured")

doc.add_heading('H.5. Launch Checklist', level=2)
checklist = [
    "Homepage loads correctly",
    "College search works with all filters",
    "College pages display all information",
    "Course pages work correctly",
    "Admin panel accessible",
    "Content can be edited without code changes",
    "Enquiries reach admin panel",
    "Authentication works (login/register/OTP)",
    "Mobile version works on all devices",
    "SEO basics configured (title, description, sitemap)",
    "HTTPS enabled and working",
    "Database backup automated",
    "Security testing completed",
    "All major APIs tested",
    "Error handling works (no technical errors shown)",
    "Analytics configured (Google Analytics)",
    "Sitemap available at /sitemap.xml",
    "Production deployment stable",
    "Monitoring alerts configured",
    "Backup restoration tested",
]
for item in checklist:
    add_bullet(f"☐ {item}")

doc.add_page_break()

# ============================================
# SECTION I: FUTURE EXPANSION
# ============================================
doc.add_heading('SECTION I: FUTURE EXPANSION ROADMAP', level=1)

doc.add_heading('I.1. Short Term (6-12 months after launch)', level=2)
add_bullet("Online counselling system")
add_bullet("Application tracking")
add_bullet("Student document management")
add_bullet("Push notifications")
add_bullet("Multilingual support (Hindi, Kannada, Tamil)")
add_bullet("Advanced analytics dashboard")

doc.add_heading('I.2. Medium Term (12-24 months)', level=2)
add_bullet("Mobile application (React Native)")
add_bullet("Online payments (Stripe/Razorpay)")
add_bullet("Premium memberships")
add_bullet("College advertising")
add_bullet("Sponsored listings")
add_bullet("College CRM")
add_bullet("Partner college dashboard")

doc.add_heading('I.3. Long Term (24+ months)', level=2)
add_bullet("Student application dashboard")
add_bullet("Scholarship application tracking")
add_bullet("AI career counselling")
add_bullet("Personalized course recommendations")
add_bullet("Live counselling sessions")
add_bullet("Peer-to-peer student community")
add_bullet("International colleges")

doc.add_page_break()

# ============================================
# SECTION J: APPENDICES
# ============================================
doc.add_heading('SECTION J: APPENDICES', level=1)

doc.add_heading('J.1. Complete Environment Variables List', level=2)
envs_full = [
    ("DATABASE_URL", "PostgreSQL connection string", "Required"),
    ("REDIS_URL", "Redis connection string", "Required"),
    ("JWT_SECRET", "Secret key for JWT tokens", "Required"),
    ("JWT_ALGORITHM", "HS256", "Required"),
    ("JWT_ACCESS_EXPIRY", "900 (15 minutes)", "Required"),
    ("JWT_REFRESH_EXPIRY", "604800 (7 days)", "Required"),
    ("AI_API_KEY", "OpenAI/Cohere API key", "Required"),
    ("AI_MODEL", "gpt-4 / gpt-3.5-turbo", "Required"),
    ("EMAIL_PROVIDER", "sendgrid / ses", "Required"),
    ("EMAIL_API_KEY", "SendGrid/SES API key", "Required"),
    ("EMAIL_FROM", "noreply@padhaanewala.in", "Required"),
    ("SMS_PROVIDER", "msg91 / twilio", "Required"),
    ("SMS_API_KEY", "MSG91/Twilio API key", "Required"),
    ("WHATSAPP_API_KEY", "Twilio/WATI API key", "Optional"),
    ("STORAGE_PROVIDER", "s3 / r2", "Required"),
    ("STORAGE_ACCESS_KEY", "S3/R2 access key", "Required"),
    ("STORAGE_SECRET", "S3/R2 secret key", "Required"),
    ("STORAGE_BUCKET", "padhaanewala-media", "Required"),
    ("STORAGE_REGION", "ap-south-1", "Required"),
    ("SENTRY_DSN", "Sentry error tracking DSN", "Recommended"),
    ("GOOGLE_ANALYTICS_ID", "G-XXXXXXXXXX", "Recommended"),
    ("GOOGLE_SEARCH_CONSOLE", "Verification code", "Recommended"),
    ("GOOGLE_MAPS_API_KEY", "Google Maps API key", "Recommended"),
    ("NEXT_PUBLIC_API_URL", "https://api.padhaanewala.in", "Required"),
    ("NEXT_PUBLIC_SITE_URL", "https://padhaanewala.in", "Required"),
    ("CORS_ORIGINS", "https://padhaanewala.in", "Required"),
    ("RATE_LIMIT_PER_MINUTE", "100", "Required"),
    ("CELERY_BROKER_URL", "redis://localhost:6379/1", "Required"),
    ("CELERY_RESULT_BACKEND", "redis://localhost:6379/2", "Required"),
]
add_table(["Variable", "Value", "Status"], envs_full)

doc.add_heading('J.2. Third-party Services', level=2)
services = [
    ("Cloudflare", "CDN + DDoS Protection + WAF", "Free / Pro $20/mo"),
    ("AWS S3 / Cloudflare R2", "Object Storage", "~$5/mo"),
    ("SendGrid", "Email Service", "Free (100/day) / Pro $20/mo"),
    ("MSG91", "SMS Service (India)", "~₹0.20/SMS"),
    ("OpenAI", "AI/LLM API", "Pay per use"),
    ("pgvector", "Vector Search", "Free (PostgreSQL extension)"),
    ("Sentry", "Error Tracking", "Free tier / Team $26/mo"),
    ("Google Analytics", "Website Analytics", "Free"),
    ("Google Maps", "Maps API", "Free tier / Pay per use"),
    ("GitHub", "Version Control", "Free"),
    ("Docker", "Containerization", "Free"),
]
add_table(["Service", "Purpose", "Cost"], services)

doc.add_heading('J.3. Cost Estimation (Monthly)', level=2)
costs = [
    ("Cloud Server (4 vCPU, 8GB)", "AWS/DigitalOcean", "$40-80"),
    ("Cloudflare Pro", "CDN + WAF", "$20"),
    ("AWS S3 / R2", "Storage (100GB)", "$5"),
    ("SendGrid", "Email (1000/day)", "$20"),
    ("MSG91", "SMS (1000/month)", "$25"),
    ("OpenAI API", "AI (moderate use)", "$50-100"),
    ("Domain", "padhaanewala.in", "$12/year"),
    ("SSL Certificate", "Cloudflare Free", "$0"),
    ("Sentry", "Error Tracking", "$0-26"),
    ("TOTAL (Minimum)", "", "$160-280/month"),
    ("TOTAL (Recommended)", "", "$250-400/month"),
]
add_table(["Item", "Provider", "Cost"], costs)

# ============================================
# ASSIGNED ROLE STATEMENT
# ============================================
doc.add_page_break()
doc.add_heading('ASSIGNED ROLE STATEMENT', level=1)
add_body("As the assigned Website Engineering & Product Development Team of Padhaanewala Edutech Services, Bengaluru – 560100, the engineering team is responsible for translating this specification into a secure, scalable, maintainable and production-ready education technology platform.")
add_body("Any architectural deviation that materially affects scalability, SEO, security, data ownership, admin control or future expansion should be discussed and approved before implementation.")
add_body("This Version 2.0 specification includes the original PDF requirements (Sections 1-69), RAG integration (Section B), Architecture Improvements (Section C), Database Design (Section D), API Design (Section E), Implementation Plan (Section F), Testing Plan (Section G), Deployment Guide (Section H), Future Expansion (Section I), and Appendices (Section J).")

doc.add_paragraph()
add_title("PADHAANEWALA EDUTECH SERVICES", size=14)
add_title("BENGALURU – 560100", size=12, bold=False)
add_title("End of Master Website Development Specification Version 2.0", size=12, bold=False)

# ============================================
# SAVE DOCUMENT
# ============================================
output_path = r"D:\code\Clients\Padhaanewala\Padhaanewala_Complete_Master_Specification.docx"
doc.save(output_path)
print(f"Document saved to: {output_path}")
