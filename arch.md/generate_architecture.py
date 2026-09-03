from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Consolas'
style.font.size = Pt(9)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.0

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.color.rgb = RGBColor(0, 51, 102)
    hs.font.bold = True
    if level == 1:
        hs.font.size = Pt(18)
        hs.paragraph_format.space_before = Pt(24)
        hs.paragraph_format.space_after = Pt(12)
    elif level == 2:
        hs.font.size = Pt(15)
        hs.paragraph_format.space_before = Pt(18)
        hs.paragraph_format.space_after = Pt(8)
    else:
        hs.font.size = Pt(13)
        hs.paragraph_format.space_before = Pt(12)
        hs.paragraph_format.space_after = Pt(6)

def add_title(text, size=24, bold=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0, 51, 102)
    return p

def add_body(text):
    doc.add_paragraph(text)

def add_code(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    return p

def add_bullet(text):
    doc.add_paragraph(text, style='List Bullet')

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            table.rows[ri + 1].cells[ci].text = str(val)
    doc.add_paragraph()
    return table

# ============================================
# COVER PAGE
# ============================================
for _ in range(3):
    doc.add_paragraph()

add_title("PADHAANEWALA EDUTECH SERVICES")
add_title("BENGALURU - 560100", size=16, bold=False)
doc.add_paragraph()
add_title("COMPLETE SYSTEM ARCHITECTURE", size=26)
doc.add_paragraph()
add_title("Padhaanewala Education Technology Platform", size=14, bold=False)
add_title("padhaanewala.in", size=12, bold=False)
add_title("Version: 1.0", size=12, bold=False)
add_title("Date: September 2026", size=12, bold=False)

doc.add_page_break()

# ============================================
# 1. OVERALL SYSTEM ARCHITECTURE
# ============================================
doc.add_heading('1. OVERALL SYSTEM ARCHITECTURE', level=1)
add_body("The complete Padhaanewala platform architecture showing all components and how they connect:")

add_code("""
                    ┌─────────────────────────────────────────────────────────┐
                    │                         INTERNET                        │
                    └───────────────────────────┬─────────────────────────────┘
                                                │ HTTPS (443)
                                                ▼
                    ┌─────────────────────────────────────────────────────────┐
                    │                       CLOUDFLARE                        │
                    │                 (CDN + WAF + DDoS Protection)           │
                    └───────────────────────────┬─────────────────────────────┘
                                                │
                                                ▼
                    ┌─────────────────────────────────────────────────────────┐
                    │                  LOAD BALANCER (Nginx)                  │
                    │                  SSL Termination + Routing               │
                    └───────────────────────────┬─────────────────────────────┘
                                                │
                 ┌──────────────────────────────┼──────────────────────────────┐
                 │                              │                              │
                 ▼                              ▼                              ▼
        ┌──────────────────┐          ┌──────────────────┐          ┌──────────────────┐
        │   NEXT.JS APP 1  │          │   NEXT.JS APP 2  │          │   NEXT.JS APP 3  │
        │  (React + TS)    │          │  (React + TS)    │          │  (React + TS)    │
        │  Public Website  │          │                  │          │                  │
        │  Student Portal  │          │   Admin Portal   │          │   CRM Portal     │
        │  SEO Pages       │          │   CMS            │          │   Counsellor     │
        └────────┬─────────┘          └────────┬─────────┘          └────────┬─────────┘
                 │                             │                             │
                 └─────────────────────────────┼─────────────────────────────┘
                                               │ REST API (HTTPS)
                                               ▼
                    ┌─────────────────────────────────────────────────────────┐
                    │                  LOAD BALANCER (Nginx)                  │
                    │                    API Gateway                          │
                    │              Route: /api/v1/* to FastAPI                │
                    └───────────────────────────┬─────────────────────────────┘
                                                │
                 ┌──────────────────────────────┼──────────────────────────────┐
                 │                              │                              │
                 ▼                              ▼                              ▼
        ┌──────────────────┐          ┌──────────────────┐          ┌──────────────────┐
        │   FASTAPI APP 1  │          │   FASTAPI APP 2  │          │   FASTAPI APP 3  │
        │                  │          │                  │          │                  │
        │  Bussiness Logic │          │   REST APIs      │          │   AI Services    │
        │  Auth/Users      │          │   Colleges       │          │   Predictor      │
        │  Search          │          │   Courses        │          │   Assistant      │
        │  CRM/Leads       │          │   Mock Tests     │          │   RAG Pipeline   │
        └────────┬─────────┘          └────────┬─────────┘          └────────┬─────────┘
                 │                             │                             │
                 └─────────────────────────────┼─────────────────────────────┘
                                               │
        ┌──────────────────────────────┬───────┼──────────┬──────────────────────────┐
        │                              │               │                          │
        ▼                              ▼               ▼                          ▼
┌───────────────┐            ┌───────────────┐  ┌───────────────┐          ┌───────────────┐
│   PostgreSQL  │            │     Redis     │  │ Celery Worker │          │ Object Storage│
│               │            │               │  │               │          │   S3 / R2     │
│  Main DB      │            │  Cache        │  │ Background    │          │               │
│  All Tables   │            │  Sessions     │  │ Jobs:         │          │  Images       │
│  pgvector     │            │  Rate Limits  │  │  Embeddings   │          │  Uploads      │
│  Search       │            │  API Cache    │  │  CSV Import   │          │  Documents    │
│               │            │  Queues       │  │  Emails       │          │  Media        │
└───────────────┘            └───────────────┘  │  SMS          │          └───────────────┘
                                                │  Images       │
                                                └───────────────┘
                                                      │
                                                      ▼
                                     ┌──────────────────────────────┐
                                     │      EXTERNAL SERVICES       │
                                     │                              │
                                     │  OpenAI / Cohere (AI)       │
                                     │  SendGrid (Email)           │
                                     │  MSG91 (SMS)                │
                                     │  Twilio (WhatsApp)          │
                                     │  Google Maps                │
                                     │  Google Analytics           │
                                     │  Sentry (Monitoring)        │
                                     │  Datadog (Performance)      │
                                     └──────────────────────────────┘

        MONITORING & OBSERVABILITY LAYER
        ┌─────────────────────────────────────────────────────────┐
        │  Sentry (Errors)  │  Datadog (Perf)  │  UptimeRobot     │
        │  Prometheus (Metrics) │ Grafana (Dashboards) │ ELK (Logs) │
        └─────────────────────────────────────────────────────────┘

        DEVOPS / CI-CD LAYER
        ┌─────────────────────────────────────────────────────────┐
        │  Git → GitHub → GitHub Actions → Docker → Production    │
        │  Docker Compose │ Docker Swarm / Kubernetes │ Terraform │
        └─────────────────────────────────────────────────────────┘
""")

doc.add_page_break()

# ============================================
# 2. FRONTEND ARCHITECTURE
# ============================================
doc.add_heading('2. FRONTEND ARCHITECTURE (NEXT.JS)', level=1)
add_body("Frontend is built with Next.js (App Router), React, TypeScript, and Tailwind CSS.")

add_code("""
                    NEXT.JS APP (App Router)
    ┌─────────────────────────────────────────────────────────────┐
    │                                                             │
    │  PUBLIC ROUTES                STUDENT ROUTES                │
    │  ┌──────────────┐             ┌──────────────────┐          │
    │  │ /            │ Home        │ /dashboard       │ Profile │
    │  │ /colleges    │ Search      │ /dashboard/me    │ Info    │
    │  │ /college/xyz │ Detail      │ /dashboard/me/   │ Saved   │
    │  │ /courses     │ List        │   colleges       │ Colleges│
    │  │ /course/bhms │ Detail      │ /dashboard/me/   │ Tests   │
    │  │ /scholarships│ List        │   tests          │ History │
    │  │ /exams       │ List        │ /compare         │ Compare │
    │  │ /mock-tests  │ List        │ /predictor       │ AI      │
    │  │ /ai          │ Assistant   │ /scholarships    │ Filters │
    │  │ /blog        │ List        │ └──────────────────┘        │
    │  │ /blog/slug   │ Article                                     │
    │  │ /about /contact /legal                                    │
    │  └──────────────┘            ADMIN ROUTES                   │
    │                              ┌──────────────────┐           │
    │  AUTH ROUTES                 │ /admin           │ Dashboard│
    │  ┌──────────────┐            │ /admin/colleges  │ Manage   │
    │  │ /login       │            │ /admin/courses   │ Manage   │
    │  │ /register    │            │ /admin/scholarships         │
    │  │ /forgot-     │            │ /admin/exams     │ Manage   │
    │  │   password   │            │ /admin/tests     │ Manage   │
    │  └──────────────┘            │ /admin/leads     │ Manage   │
    │                              │ /admin/counsellors        │
    │  COUNSELLOR ROUTES           │ /admin/reviews   │ Manage   │
    │  ┌──────────────────┐        │ /admin/blogs     │ Manage   │
    │  │ /counsellor      │ Leads  │ /admin/media     │ Manage   │
    │  │ /counsellor/me   │ Notes  │ /admin/seo       │ Manage   │
    │  │ /counsellor/me/  │ Follow │ /admin/settings  │ Config   │
    │  │   leads          │ up     │ └──────────────────┘        │
    │  └──────────────────┘                                      │
    │                                                             │
    └─────────────────────────────────────────────────────────────┘

    COMPONENT LAYERS
    ┌─────────────────────────────────────────────────────────────┐
    │  UI Components (Button, Card, Input, Modal, Table)          │
    │  Feature Components (CollegeCard, CourseCard, SearchBar)    │
    │  Layout Components (Header, Footer, Nav, Sidebar)           │
    │  Page Components (HomePage, CollegePage, AdminPage)         │
    │  HOCs / Providers (AuthProvider, ThemeProvider, Query)      │
    └─────────────────────────────────────────────────────────────┘

    STATE MANAGEMENT
    ┌─────────────────────────────────────────────────────────────┐
    │  Server State: TanStack Query (React Query)                 │
    │  Client State: Zustand / Context API                        │
    │  Form State: React Hook Form + Zod Validation               │
    │  Routing: Next.js App Router                                │
    │  Styling: Tailwind CSS                                      │
    │  Data Fetch: SWR / TanStack Query                           │
    └─────────────────────────────────────────────────────────────┘
""")

doc.add_heading('2.1. Frontend File Structure', level=2)
add_code("""
frontend/
├── app/
│   ├── (public)/                  # Public facing routes
│   │   ├── page.tsx               # Homepage
│   │   ├── colleges/
│   │   │   ├── page.tsx           # College search
│   │   │   └── [slug]/page.tsx    # College detail
│   │   ├── courses/
│   │   │   ├── page.tsx
│   │   │   └── [slug]/page.tsx
│   │   ├── scholarships/
│   │   │   ├── page.tsx
│   │   │   └── [slug]/page.tsx
│   │   ├── exams/
│   │   │   ├── page.tsx
│   │   │   └── [slug]/page.tsx
│   │   ├── mock-tests/
│   │   │   ├── page.tsx
│   │   │   ├── [id]/page.tsx      # Test taking
│   │   │   └── [id]/result/page.tsx
│   │   ├── predictor/page.tsx     # AI College Predictor
│   │   ├── ai/page.tsx            # AI Education Assistant
│   │   ├── compare/page.tsx       # College Comparison
│   │   ├── blog/
│   │   │   ├── page.tsx
│   │   │   └── [slug]/page.tsx
│   │   ├── about/page.tsx
│   │   ├── contact/page.tsx
│   │   └── legal/[page]/page.tsx  # Privacy, Terms, etc
│   │
│   ├── (auth)/
│   │   ├── login/page.tsx
│   │   ├── register/page.tsx
│   │   └── forgot-password/page.tsx
│   │
│   ├── (student)/
│   │   └── dashboard/
│   │       ├── page.tsx
│   │       ├── profile/page.tsx
│   │       ├── colleges/page.tsx      # Saved colleges
│   │       ├── tests/page.tsx         # Test history
│   │       └── enquiries/page.tsx
│   │
│   ├── (counsellor)/
│   │   └── counsellor/
│   │       ├── page.tsx
│   │       └── leads/[id]/page.tsx
│   │
│   ├── (admin)/
│   │   └── admin/
│   │       ├── page.tsx               # Dashboard
│   │       ├── colleges/page.tsx
│   │       ├── courses/page.tsx
│   │       ├── scholarships/page.tsx
│   │       ├── exams/page.tsx
│   │       ├── tests/page.tsx
│   │       ├── students/page.tsx
│   │       ├── reviews/page.tsx
│   │       ├── leads/page.tsx
│   │       ├── counsellors/page.tsx
│   │       ├── blogs/page.tsx
│   │       ├── media/page.tsx
│   │       ├── seo/page.tsx
│   │       └── settings/page.tsx
│   │
│   ├── api/                          # API routes (proxy)
│   │   └── [...path]/route.ts
│   │
│   ├── layout.tsx
│   └── globals.css
│
├── components/
│   ├── ui/                          # Reusable UI components
│   │   ├── Button.tsx
│   │   ├── Card.tsx
│   │   ├── Input.tsx
│   │   ├── Modal.tsx
│   │   ├── Dropdown.tsx
│   │   └── ...
│   ├── layout/
│   │   ├── Header.tsx
│   │   ├── Footer.tsx
│   │   └── Navbar.tsx
│   ├── features/
│   │   ├── colleges/
│   │   │   ├── CollegeCard.tsx
│   │   │   ├── CollegeFilters.tsx
│   │   │   └── CollegeComparison.tsx
│   │   ├── search/
│   │   │   ├── SearchBar.tsx
│   │   │   └── SearchResults.tsx
│   │   ├── predictor/
│   │   │   └── PredictorForm.tsx
│   │   └── ai/
│   │       └── ChatWidget.tsx
│   ├── forms/
│   │   ├── EnquiryForm.tsx
│   │   └── LoginForm.tsx
│   └── schema/                      # Schema.org markup
│       ├── CollegeSchema.tsx
│       ├── CourseSchema.tsx
│       └── FAQSchema.tsx
│
├── hooks/
│   ├── useAuth.ts
│   ├── useSearch.ts
│   ├── useDebounce.ts
│   └── useInfiniteScroll.ts
│
├── lib/
│   ├── api-client.ts               # Axios/Fetch wrapper
│   ├── auth.ts                     # JWT handling
│   ├── utils.ts                    # Helper functions
│   └── constants.ts
│
├── types/
│   ├── college.ts
│   ├── course.ts
│   ├── user.ts
│   └── index.ts
│
├── middleware.ts                    # Route protection
├── next.config.mjs
├── tailwind.config.ts
└── tsconfig.json
""")

doc.add_page_break()

# ============================================
# 3. BACKEND ARCHITECTURE
# ============================================
doc.add_heading('3. BACKEND ARCHITECTURE (FASTAPI)', level=1)
add_body("Backend is built with Python FastAPI, following a modular service-oriented architecture with separation of concerns.")

add_code("""
                    FASTAPI BACKEND
    ┌─────────────────────────────────────────────────────────────┐
    │                                                             │
    │  API LAYER (Routers)                                        │
    │  ├── /api/v1/auth          → AuthRouter                     │
    │  ├── /api/v1/users         → UserRouter                     │
    │  ├── /api/v1/colleges      → CollegeRouter                  │
    │  ├── /api/v1/courses       → CourseRouter                   │
    │  ├── /api/v1/search        → SearchRouter                   │
    │  ├── /api/v1/predictor     → PredictorRouter                │
    │  ├── /api/v1/ai            → AIRouter                       │
    │  ├── /api/v1/mock-tests    → MockTestRouter                 │
    │  ├── /api/v1/scholarships  → ScholarshipRouter              │
    │  ├── /api/v1/exams         → ExamRouter                     │
    │  ├── /api/v1/reviews       → ReviewRouter                   │
    │  ├── /api/v1/blog          → BlogRouter                     │
    │  ├── /api/v1/enquiries     → EnquiryRouter                  │
    │  └── /api/v1/admin         → AdminRouter                    │
    │                                                             │
    │  SERVICE LAYER (Business Logic)                             │
    │  ├── auth_service.py        Auth/OTP/JWT                    │
    │  ├── college_service.py     College operations              │
    │  ├── course_service.py      Course operations               │
    │  ├── search_service.py      Search + filters                │
    │  ├── predictor_service.py   AI College Predictor            │
    │  ├── ai_service.py          AI Assistant (RAG)              │
    │  ├── mock_test_service.py   Test engine                     │
    │  ├── scholarship_service.py Scholarship ops                  │
    │  ├── lead_service.py        Lead/CRM operations             │
    │  ├── content_service.py     CMS operations                  │
    │  └── analytics_service.py   Analytics                       │
    │                                                             │
    │  REPOSITORY LAYER (Data Access)                             │
    │  ├── college_repo.py        SQL queries                     │
    │  ├── course_repo.py         SQL queries                     │
    │  ├── user_repo.py           SQL queries                     │
    │  └── ... (one per entity)                                   │
    │                                                             │
    │  CORE LAYER                                                  │
    │  ├── config.py              Settings/env vars               │
    │  ├── security.py            JWT, bcrypt, OTP                │
    │  ├── permissions.py         RBAC roles                      │
    │  ├── database.py            SQLAlchemy engine               │
    │  ├── redis.py               Cache connection                │
    │  ├── cache.py               Caching utilities               │
    │  └── logging.py             Structured logging              │
    │                                                             │
    │  AI / RAG LAYER                                              │
    │  ├── embedding_service.py   Embedding generation            │
    │  ├── vector_search.py       pgvector search                 │
    │  ├── rag_pipeline.py        RAG orchestration               │
    │  ├── predictor_engine.py    Predictor logic                 │
    │  └── prompt_templates.py    LLM prompts                     │
    │                                                             │
    │  CELERY TASKS                                                │
    │  ├── tasks/
    │  │   ├── import_tasks.py    CSV import                      │
    │  │   ├── embedding_tasks.py Embedding generation            │
    │  │   ├── email_tasks.py     Email sending                   │
    │  │   ├── sms_tasks.py       SMS sending                     │
    │  │   ├── image_tasks.py     Image optimization              │
    │  │   ├── sitemap_tasks.py   Sitemap generation              │
    │  │   └── analytics_tasks.py Analytics aggregation           │
    │                                                             │
    │  SCHEMAS (Pydantic)                                          │
    │  ├── auth_schemas.py        Login/Register/Otp              │
    │  ├── college_schemas.py     College DTOs                    │
    │  ├── course_schemas.py      Course DTOs                     │
    │  ├── search_schemas.py      Search DTOs                     │
    │  └── ... (one per module)                                   │
    │                                                             │
    │  MIDDLEWARE (Order Matters)                                  │
    │  ├── 1. CORS                Allow padhaanewala.in          │
    │  ├── 2. Rate Limiter        IP + User limits                │
    │  ├── 3. Security Headers    Helmet-like                     │
    │  ├── 4. Request Logger      Log all requests                │
    │  ├── 5. Error Handler       Global exception handling       │
    │  └── 6. Auth Dependency     JWT verification                │
    │                                                             │
    └─────────────────────────────────────────────────────────────┘
""")

doc.add_page_break()

# ============================================
# 4. DATABASE ARCHITECTURE
# ============================================
doc.add_heading('4. DATABASE ARCHITECTURE (POSTGRESQL)', level=1)

doc.add_heading('4.1. Entity Relationship Overview', level=2)
add_code("""
                        USERS
   ┌──────────────────────┼───────────────────────┐
   ▼                      ▼                       ▼
STUDENT_PROFILES      ADMINS               COUNSELLORS
   │                                            │
   ├── education                                 ├── max_leads
   ├── course_interest                           └── availability
   ├── preferred_state
   └── budget

   ┌──────────────────┬──────────────────┐
   ▼                  ▼                  ▼
SAVED_COLLEGES    TEST_ATTEMPTS     ENQUIRIES
   │                  │                  │
   ▼                  ▼                  ▼
COLLEGES          MOCK_TESTS       LEAD_NOTES
   │                  │                  │
   ├── university ────┤                  ▼
   ├── state          │            COUNSELLORS (assigned)
   ├── city           ▼
   ├── latitude   QUESTIONS
   └── longitude      │
                      ▼
                 ANSWERS (or in JSON options)

   COLLEGES ─────── COLLEGE_COURSES ─────── COURSES
        │                  │                     │
        │                  ▼                     ├── degree
        │                  ├── FEES              ├── duration
        │                  └── ADMISSIONS        ├── eligibility
        │                                         └── entrance_exam
        │
        ├── REVIEWS
        ├── GALLERY (MEDIA)
        └── FAQS

   SCHOLARSHIPS ──── (filters: state, course, income)

   EXAMS ──── EXAM_DATES

   BLOGS ──── CATEGORIES ──── SEO_METADATA

   ALL ENTITIES ──── MEDIA (generic storage)
        │
        └── DOCUMENT_EMBEDDINGS (RAG vector storage)
""")

doc.add_heading('4.2. Complete Table Structure', level=2)

tables = [
    ("users", "User accounts (students, admins, counsellors)", "id, email, phone, password_hash, role, is_active, created_at, updated_at"),
    ("student_profiles", "Student additional info", "id, user_id, name, education, course_interest, preferred_state, preferred_city, budget"),
    ("admins", "Admin accounts", "id, user_id, role, permissions"),
    ("counsellors", "Counsellor accounts", "id, user_id, name, phone, email, max_leads"),
    ("colleges", "College master data", "id, name, slug, official_name, type, is_government, university_id, state, district, city, address, pincode, website, email, phone, established_year, accreditation, recognition, description, latitude, longitude, verification_status"),
    ("courses", "Course master data", "id, name, slug, degree, duration, description, eligibility, entrance_exam, admission_procedure, career_info, seo_title, seo_description"),
    ("college_courses", "Many-to-many colleges↔courses", "id, college_id, course_id, fees_id, admission_info, seats"),
    ("universities", "University master data", "id, name, slug, state, website, description"),
    ("locations", "Location reference data", "id, state, district, city, latitude, longitude"),
    ("fees", "Structured fee data", "id, college_course_id, tuition_fee, hostel_fee, exam_fee, other_charges, total_approximate, fee_period, is_approximate"),
    ("admissions", "Admission info per course", "id, college_course_id, eligibility, entrance_exam, cutoff, admission_process, important_dates"),
    ("scholarships", "Scholarship master data", "id, name, slug, provider, is_government, eligibility, state, course, income_criteria, amount, deadline, documents, application_procedure, official_link, status"),
    ("exams", "Exam master data", "id, name, slug, conducting_authority, eligibility, application_start, application_deadline, exam_date, admit_card_date, result_date, official_website"),
    ("mock_tests", "Mock test configuration", "id, name, exam_id, subject, difficulty, total_questions, time_limit, instructions, status"),
    ("questions", "Test questions", "id, mock_test_id, question_text, question_type, options(JSON), correct_answer, explanation, marks"),
    ("test_attempts", "Student test attempts", "id, user_id, mock_test_id, answers(JSON), score, percentage, correct, incorrect, unattempted, time_taken, topic_wise_performance(JSON), rank, percentile"),
    ("reviews", "Student reviews", "id, user_id, college_id, course, year, rating, review_text, images(JSON), status, moderation_notes"),
    ("blogs", "Blog articles", "id, title, slug, content, featured_image, category_id, author_id, meta_title, meta_description, canonical_url, status, published_at"),
    ("categories", "Content categories", "id, name, slug, description, parent_id, type"),
    ("enquiries", "Admission enquiries / leads", "id, user_id, name, phone, email, course, preferred_college, state, city, qualification, message, source, utm_source, utm_medium, utm_campaign, utm_content, status, assigned_counsellor_id"),
    ("lead_notes", "Counsellor notes on leads", "id, enquiry_id, counsellor_id, note, created_at"),
    ("notifications", "User notifications", "id, user_id, title, message, type, is_read, created_at"),
    ("saved_colleges", "Student saved colleges", "id, user_id, college_id, created_at"),
    ("faqs", "FAQs for any entity", "id, entity_type, entity_id, question, answer, order, status"),
    ("media", "Media files", "id, filename, original_name, mime_type, size, url, alt_text, entity_type, entity_id, uploaded_by"),
    ("seo_metadata", "Per-page SEO data", "id, entity_type, entity_id, title, description, canonical_url, og_title, og_description, og_image, schema_markup"),
    ("audit_logs", "Security/action logs", "id, user_id, action, entity_type, entity_id, old_value, new_value, ip_address, user_agent"),
    ("document_embeddings", "RAG vector storage", "id, entity_type, entity_id, chunk_text, embedding(VECTOR), metadata(JSONB)"),
]
add_table(["Table", "Purpose", "Key Columns"], tables)

doc.add_heading('4.3. Relationships (Foreign Keys)', level=2)
add_code("""
users.id                ← student_profiles.user_id (1:1)
users.id                ← admins.user_id (1:1)
users.id                ← counsellors.user_id (1:1)
users.id                ← saved_colleges.user_id (1:many)
users.id                ← reviews.user_id (1:many)
users.id                ← test_attempts.user_id (1:many)
users.id                ← notifications.user_id (1:many)
users.id                ← enquiries.user_id (1:many)

colleges.id             ← college_courses.college_id (1:many)
courses.id              ← college_courses.course_id (1:many)
college_courses.id      ← fees.college_course_id (1:1)
college_courses.id      ← admissions.college_course_id (1:1)

universities.id         ← colleges.university_id (1:many)
colleges.id             ← reviews.college_id (1:many)
colleges.id             ← saved_colleges.college_id (1:many)

mock_tests.id           ← questions.mock_test_id (1:many)
mock_tests.id           ← test_attempts.mock_test_id (1:many)

enquiries.id            ← lead_notes.enquiry_id (1:many)
counsellors.id          ← enquiries.assigned_counsellor_id (1:many)

blogs.category_id       ← categories.id
blogs.author_id         ← users.id
""")

doc.add_heading('4.4. Indexes', level=2)
add_code("""
-- Exclude FROM indexing for query performance:
CREATE INDEX idx_colleges_slug     ON colleges(slug);
CREATE INDEX idx_colleges_state    ON colleges(state);
CREATE INDEX idx_colleges_city     ON colleges(city);
CREATE INDEX idx_colleges_type     ON colleges(type);
CREATE INDEX idx_colleges_status   ON colleges(status);

CREATE INDEX idx_courses_slug      ON courses(slug);
CREATE INDEX idx_courses_status    ON courses(status);

CREATE INDEX idx_college_courses_college ON college_courses(college_id);
CREATE INDEX idx_college_courses_course  ON college_courses(course_id);

CREATE INDEX idx_scholarships_state ON scholarships(state);
CREATE INDEX idx_scholarships_deadline ON scholarships(deadline);

CREATE INDEX idx_exams_date ON exams(exam_date);
CREATE INDEX idx_reviews_college ON reviews(college_id, status);

CREATE INDEX idx_blog_slug ON blogs(slug, status);
CREATE INDEX idx_enquiries_status ON enquiries(status, assigned_counsellor_id);

CREATE INDEX idx_users_email ON users(email);
CREATE INDEX idx_users_phone ON users(phone);

-- pgvector index for RAG:
CREATE INDEX idx_embeddings_vector 
  ON document_embeddings 
  USING ivfflat (embedding vector_cosine_ops) 
  WITH (lists = 100);

-- JSONB indexes for quick metadata queries:
CREATE INDEX idx_embeddings_entity 
  ON document_embeddings(entity_type, entity_id);
""")

doc.add_page_break()

# ============================================
# 5. RAG / AI ARCHITECTURE
# ============================================
doc.add_heading('5. AI / RAG ARCHITECTURE', level=1)

doc.add_heading('5.1. RAG Pipeline Flow', level=2)
add_code("""
              STUDENT QUERY
                   │
                   ▼
        ┌───────────────────┐
        │  QUERY PREPROCESS │───→ Extract intent + entities
        └─────────┬─────────┘
                  │
        ┌─────────┴─────────┐
        │  STRUCTURED SEARCH│─→ PostgreSQL filters (course, state, fee...)
        └─────────┬─────────┘
                  │
        ┌─────────┴─────────┐
        │  EMBEDDING GEN    │─→ Convert query to vector (OpenAI)
        └─────────┬─────────┘
                  │
        ┌─────────┴─────────┐
        │   VECTOR SEARCH   │─→ pgvector similarity (top-K chunks)
        └─────────┬─────────┘
                  │
        ┌─────────┴─────────┐
        │ CONTEXT ASSEMBLY  │─→ Combine DB + vector results
        └─────────┬─────────┘
                  │
        ┌─────────┴─────────┐
        │   LLM PROCESSING  │─→ GPT-4 / Claude (context + query)
        └─────────┬─────────┘
                  │
        ┌─────────┴─────────┐
        │   RESPONSE +      │─→ Answer + source citations + disclaimer
        │   SOURCES         │
        └───────────────────┘
                │
                ▼
           STUDENT SEES
""")

doc.add_heading('5.2. Embedding & Vector Pipeline', level=2)
add_code("""
   DATA CHANGE (Admin adds/edits college)
              │
              ▼
      CELERY TASK: generate_embeddings
              │
      ┌───────┴────────┐
      │ CHUNK DATA     │
      └───────┬────────┘
              │
      ┌───────┴────────┐
      │ EMBEDDING API  │ (OpenAI text-embedding-3-small)
      └───────┬────────┘
              │
      ┌───────┴────────┐
      │ STORE IN       │──→ document_embeddings table
      │ pgvector       │
      └────────────────┘

   DOCUMENT_EMBEDDINGS TABLE:
   +----------------+------------------+------------------+
   | entity_type    | entity_id        | chunk_text       |
   +----------------+------------------+------------------+
   | college        | 123              | "XYZ Medical..." |
   | course         | 45               | "BHMS is a ..."  |
   | scholarship    | 89               | "Bihar Scholar..."|
   | faq            | 12               | "Q: What is BHMS?"|
   +----------------+------------------+------------------+
   | embedding (VECTOR(1536)) | metadata (JSONB)       |
   +--------------------------+--------------------------+
""")

doc.add_heading('5.3. AI Features Architecture', level=2)

doc.add_heading('AI College Predictor', level=3)
add_code("""
   INPUTS: course, exam, rank, category, state, budget, hostel, type
                    │
                    ▼
   ┌─────────────────────────────┐
   │ 1. Structured filters       │
   │    course=BHMS, state=KA    │
   │    max_fee=500000           │
   └─────────────┬───────────────┘
                 │
   ┌─────────────┴───────────────┐
   │ 2. DB query - eligible      │
   │    colleges (PostgreSQL)    │
   └─────────────┬───────────────┘
                 │
   ┌─────────────┴───────────────┐
   │ 3. Embedding of preferences │
   └─────────────┬───────────────┘
                 │
   ┌─────────────┴───────────────┐
   │ 4. Vector search (pgvector) │
   └─────────────┬───────────────┘
                 │
   ┌─────────────┴───────────────┐
   │ 5. Combine + rank (LLM)     │
   └─────────────┬───────────────┘
                 │
   ┌─────────────┴───────────────┐
   │ 6. Categorize:              │
   │    Highly Suitable          │
   │    Possible                 │
   │    Reach                    │
   └─────────────┬───────────────┘
                 │
                 ▼
   Output + "Results are estimates, not guaranteed admissions"
""")

doc.add_heading('AI Education Assistant', level=3)
add_code("""
   STUDENT QUESTION: "What is the difference between BAMS and BHMS?"
                    │
                    ▼
   ┌─────────────────────────────┐
   │ Intent: Course Comparison   │
   │ Entities: BAMS, BHMS        │
   └─────────────┬───────────────┘
                 │
   ┌─────────────┴───────────────┐
   │ DB: fetch BAMS + BHMS data  │
   │    from courses table       │
   └─────────────┬───────────────┘
                 │
   ┌─────────────┴───────────────┐
   │ Vector: find similar FAQs   │
   │    from embeddings          │
   └─────────────┬───────────────┘
                 │
   ┌─────────────┴───────────────┐
   │ Assemble context            │
   └─────────────┬───────────────┘
                 │
   ┌─────────────┴───────────────┐
   │ LLM generates answer        │
   └─────────────┬───────────────┘
                 │
                 ▼
   Answer + "Based on data from colleges offering these courses."
   Question flow ends with disclaimer:
   "This information may change. Please verify with institutions."
""")

doc.add_heading('5.4. AI Guardrails', level=2)
add_code("""
╔══════════════════════════════════════════════════════════════╗
║  GUARDRAIL RULES                                            ║
╠══════════════════════════════════════════════════════════════╣
║  RULE 1: NEVER invent college names, fees, or facts          ║
║  RULE 2: Only answer from retrieved database data             ║
║  RULE 3: Always cite sources for factual claims               ║
║  RULE 4: State uncertainty when data is not verified          ║
║  RULE 5: Never promise admission or guaranteed outcomes       ║
║  RULE 6: Never present regulatory info as permanent fixed     ║
║  RULE 7: Include verification disclaimer in every response    ║
║  RULE 8: If data not found, say so - do not guess            ║
╚══════════════════════════════════════════════════════════════╝
""")

doc.add_page_break()

# ============================================
# 6. SEARCH ARCHITECTURE
# ============================================
doc.add_heading('6. SEARCH ARCHITECTURE', level=1)

doc.add_heading('6.1. Search Flow', level=2)
add_code("""
   USER: "BHMS colleges in Karnataka under 5 lakh with hostel"
                    │
                    ▼
   ┌─────────────────────────────┐
   │ NATURAL LANGUAGE PARSER     │
   │ (AI/NLP or rule-based)      │
   └─────────────┬───────────────┘
                 │
   ┌─────────────┴───────────────┐
   │ EXTRACT FILTERS:            │
   │  course    = BHMS           │
   │  state     = Karnataka      │
   │  max_fee   = 500000         │
   │  hostel    = true           │
   └─────────────┬───────────────┘
                 │
   ┌─────────────┴───────────────┐
   │ BUILD SQL QUERY             │
   │ SELECT * FROM colleges      │
   │ JOIN college_courses        │
   │ WHERE course='BHMS'         │
   │  AND state='Karnataka'      │
   │  AND fees < 500000          │
   │  AND hostel = true          │
   └─────────────┬───────────────┘
                 │
   ┌─────────────┴───────────────┐
   │ EXECUTE (PostgreSQL)        │
   │ + optional vector fallback  │
   └─────────────┬───────────────┘
                 │
   ┌─────────────┴───────────────┐
   │ SORT by relevance/rating    │
   └─────────────┬───────────────┘
                 │
                 ▼
   ┌─────────────────────────────┐
   │ PAGINATED RESULTS           │
   │ (default 20 per page)       │
   └─────────────────────────────┘
""")

doc.add_heading('6.2. Search API Parameters', level=2)
add_code("""
GET /api/v1/colleges

Query params:
    course      = "BHMS"           # filter by course
    state       = "Karnataka"      # filter by state
    district    = "Bengaluru Urban"
    city        = "Bangalore"
    college_type= "private"        # government/private
    university  = "RGUHS"
    min_fee     = 100000           # minimum fee
    max_fee     = 500000           # maximum fee
    rating      = 4                # minimum rating
    hostel      = true             # hostel required
    accredited  = true             # accredited
    page        = 1                # page number
    limit       = 20               # items per page (max 50)
    sort        = "rating"         # rating | fees | name | relevance
""")

doc.add_heading('6.3. Search Types', level=2)
registered = [
    ("Basic Search", "Keyword match on name/type", "Fast, always works"),
    ("Filter Search", "Course + State + Type + Fee...", "Precise results"),
    ("Natural Language", "Free text → filters via AI/NLP", "User-friendly"),
    ("Semantic Search", "Vector similarity (pgvector)", "Understands meaning"),
    ("Hybrid Search", "DB filters + vector similarity", "Best of both"),
]
add_table(["Search Type", "How It Works", "Use Case"], registered)

doc.add_page_break()

# ============================================
# 7. CRM / LEAD ARCHITECTURE
# ============================================
doc.add_heading('7. CRM / LEAD MANAGEMENT ARCHITECTURE', level=1)

doc.add_heading('7.1. Lead Flow', level=2)
add_code("""
   STUDENT                                         ADMIN / COUNSELLOR
      │                                                    │
      ▼                                                    │
  ┌──────────────┐                                        │
  │ ENQUIRY FORM │  (College page, Course page,           │
  │ (anywhere)   │   Predictor, Homepage, Blog)           │
  └──────┬───────┘                                        │
         │ POST /api/v1/enquiries                         │
         ▼                                                │
  ┌──────────────┐                                        │
  │  RATE LIMIT  │  (anti-spam)                           │
  └──────┬───────┘                                        │
         ▼                                                │
  ┌──────────────┐                                        │
  │ VALIDATE +   │                                        │
  │ SAVE LEAD    │  (source + UTM captured)               │
  └──────┬───────┘                                        │
         │                                                │
         ▼                                                │
  ┌──────────────┐   ┌──────────────────┐                 │
  │ LEAD CREATED │──▶│ NOTIFY ADMIN     │                 │
  │ Status: New  │   │ (Email + SMS +   │                 │
  └──────┬───────┘   │  notification)   │                 │
         │           └──────────────────┘                 │
         │                                                │
         ▼                                                ▼
  ┌────────────────────────────────────────────────────────────┐
  │                    ADMIN / COUNSELLOR VIEW                  │
  │                                                             │
  │  LEAD: LD000123   Student: Rahul (98765XXXXX)               │
  │  Course: BHMS     College: XYZ Medical                      │
  │  Source: College Page    UTM: organic                       │
  │                                                             │
  │  Status: [New ▼]   Assigned To: [Counsellor 1 ▼]           │
  │  Follow-up: 2026-09-10                                      │
  │                                                             │
  │  ┌─────────────────────────────┐                            │
  │  │ NOTES: Called Rahul, wants  │                            │
  │  │ to visit college. Sent fee  │                            │
  │  │ structure on WhatsApp.      │                            │
  │  └─────────────────────────────┘                            │
  │                                                             │
  │  STATUS PIPELINE:                                          │
  │  New → Contacted → Interested → Application Started →      │
  │        → Admission Completed / Not Interested → Closed      │
  └────────────────────────────────────────────────────────────┘
""")

doc.add_heading('7.2. Lead Statuses', level=2)
statuses = [
    ("New", "Lead just submitted, not contacted yet"),
    ("Contacted", "Counsellor contacted the student"),
    ("Interested", "Student showed interest, moving forward"),
    ("Application Started", "Student started admission application"),
    ("Admission Completed", "Student got admission"),
    ("Not Interested", "Student declined, no longer pursuing"),
    ("Closed", "Lead closed (won or lost)"),
]
add_table(["Status", "Meaning"], statuses)

doc.add_heading('7.3. Lead Source Tracking', level=2)
sources = [
    ("Homepage", "Enquiry from homepage"),
    ("College Page", "Enquiry from college detail page"),
    ("Course Page", "Enquiry from course page"),
    ("Blog", "Enquiry from blog article"),
    ("Scholarship", "Enquiry from scholarship page"),
    ("Predictor", "Enquiry from AI predictor"),
    ("WhatsApp", "Enquiry via WhatsApp CTA"),
    ("Advertisement", "Enquiry from paid ads"),
    ("Organic Search", "Enquiry from Google organic"),
]
add_table(["Source", "Where Lead Came From"], sources)

doc.add_page_break()

# ============================================
# 8. SECURITY ARCHITECTURE
# ============================================
doc.add_heading('8. SECURITY ARCHITECTURE', level=1)

doc.add_heading('8.1. Security Layers', level=2)
add_code("""
   LAYER 1: NETWORK SECURITY
   ┌──────────────────────────────────────────┐
   │ Cloudflare DDoS Protection              │
   │ Cloudflare WAF (Web Application Firewall)│
   │ SSL/TLS (HTTPS only)                    │
   └──────────────────────────────────────────┘
                    │
   LAYER 2: APPLICATION SECURITY
   ┌──────────────────────────────────────────┐
   │ Rate Limiting (Redis)                   │
   │   IP: 100 req/min                       │
   │   User: 1000 req/hour                   │
   ├──────────────────────────────────────────┤
   │ Security Headers (Helmet)               │
   │   X-Content-Type-Options                │
   │   X-Frame-Options: DENY                 │
   │   HSTS                                  │
   │   CSP                                   │
   ├──────────────────────────────────────────┤
   │ CORS: only padhaanewala.in              │
   └──────────────────────────────────────────┘
                    │
   LAYER 3: INPUT SECURITY
   ┌──────────────────────────────────────────┐
   │ Pydantic Validation (all API inputs)    │
   │ SQL Injection prevention (ORM)          │
   │ XSS protection (sanitization)           │
   │ File upload validation                  │
   │ Maximum request size limits             │
   └──────────────────────────────────────────┘
                    │
   LAYER 4: AUTHENTICATION
   ┌──────────────────────────────────────────┐
   │ JWT Access Token (15 min)               │
   │ JWT Refresh Token (7 days)              │
   │ Password hashing (bcrypt)               │
   │ OTP verification                        │
   │ Admin 2FA (TOTP)                        │
   └──────────────────────────────────────────┘
                    │
   LAYER 5: AUTHORIZATION
   ┌──────────────────────────────────────────┐
   │ RBAC Roles:                             │
   │   Super Admin → everything              │
   │   Content Admin → content only          │
   │   Counsellor → assigned leads only      │
   │   Test Admin → mock tests/questions     │
   │   SEO Admin → SEO metadata only         │
   └──────────────────────────────────────────┘
                    │
   LAYER 6: AUDIT & MONITORING
   ┌──────────────────────────────────────────┐
   │ Audit logs (who did what, when, IP)     │
   │ Login attempt logging                   │
   │ Admin action logging                    │
   │ Sentry error tracking                   │
   └──────────────────────────────────────────┘
""")

doc.add_heading('8.2. JWT Implementation', level=2)
add_code("""
TOKEN STRUCTURE:
  Header:    { "alg": "HS256", "typ": "JWT" }
  Payload:   { "sub": "user_id",
               "role": "student",
               "exp": 900,         # 15 min
               "iat": current }
  Signature: HMAC-SHA256(header.payload, JWT_SECRET)

TOKEN FLOW:
  Login → Generates access_token + refresh_token
  API calls → Send Bearer access_token in header
  Token expires → Use refresh_token to get new access_token
  Logout → Revoke refresh_token
  Password change → Revoke all tokens
""")

doc.add_heading('8.3. Rate Limiting Rules', level=2)
add_code("""
ENDPOINT                     LIMIT
─────────────────────────────────────────
Public endpoints (GET)       100 req/min/IP
Search endpoints             60 req/min/IP
AI endpoints                 20 req/min/IP
Auth (login/register/OTP)    5 req/min/IP
Enquiry submission           3 req/min/IP
Admin endpoints              300 req/min/user
Upload endpoints             30 MB/min/IP
""")

doc.add_page_break()

# ============================================
# 9. DEPLOYMENT ARCHITECTURE
# ============================================
doc.add_heading('9. DEPLOYMENT & INFRASTRUCTURE ARCHITECTURE', level=1)

doc.add_heading('9.1. Environments', level=2)
envs = [
    ("Development", "localhost:3000", "Developers", "Feature development"),
    ("Staging", "staging.padhaanewala.in", "QA + Client", "Testing, demo"),
    ("Production", "padhaanewala.in", "Public", "Live platform"),
]
add_table(["Environment", "URL", "Audience", "Purpose"], envs)

doc.add_heading('9.2. Docker Architecture', level=2)
add_code("""
                     DOCKER COMPOSE
┌─────────────────────────────────────────────────────────────┐
│  SERVICE          IMAGE              PORTS                  │
│  ─────────────────────────────────────────────────────      │
│  frontend         node:20-alpine     3000:3000              │
│  backend          python:3.12        8000:8000              │
│  db               postgres:16        5432:5432              │
│  redis            redis:7-alpine     6379:6379              │
│  celery-worker    backend-image      (internal)             │
│  celery-beat      backend-image      (internal)             │
│  nginx            nginx:alpine       80:80, 443:443         │
└─────────────────────────────────────────────────────────────┘

DOCKER VOLUMES:
  pgdata       → PostgreSQL persistent data
  redisdata    → Redis persistent data
  media        → Uploaded media backup (mirrored to S3)

DOCKER NETWORKS:
  frontend-net  → Next.js + Nginx
  internal-net  → FastAPI + Postgres + Redis + Celery
""")

doc.add_heading('9.3. CI/CD Pipeline (GitHub Actions)', level=2)
add_code("""
   GIT WORKFLOW
   ┌──────────┐   ┌──────────────┐   ┌───────────┐
   │ feature/*│──▶│   develop    │──▶│   main    │
   │ branch   │   │  (staging)   │   │(production)│
   └──────────┘   └──────────────┘   └───────────┘
        │                │                │
        ▼                ▼                ▼
   ┌────────────────────────────────────────────┐
   │ GITHUB ACTIONS PIPELINE                    │
   │   on push: feature/*                       │
   │   ├── Lint                                 │
   │   ├── Run unit tests                       │
   │   └── Build Docker images                  │
   │                                            │
   │   on merge: → develop                      │
   │   ├── Deploy to STAGING                    │
   │   └── Run integration tests                │
   │                                            │
   │   on merge: → main (tag)                   │
   │   ├── Deploy to PRODUCTION                 │
   │   └── Run smoke tests                      │
   └────────────────────────────────────────────┘
""")

doc.add_heading('9.4. Backup Architecture', level=2)
add_code("""
             DAILY BACKUP SCHEDULE
   ┌───────────────────────────────────────────────┐
   │  00:00  PostgreSQL full dump → S3 (off-site)  │
   │  06:00  WAL archive → S3                      │
   │  12:00  WAL archive → S3                      │
   │  18:00  WAL archive → S3                      │
   │  Every 15m  Redis RDB → S3                    │
   │  Weekly     Restoration test (staging)        │
   └───────────────────────────────────────────────┘

   RETENTION:
   Daily backups  → 7 days
   Weekly backups → 4 weeks
   Monthly backups→ 12 months

   OBJECT STORAGE:
   Versioning enabled
   Cross-region replication → secondary bucket
""")

doc.add_page_break()

# ============================================
# 10. MONITORING ARCHITECTURE
# ============================================
doc.add_heading('10. MONITORING & OBSERVABILITY ARCHITECTURE', level=1)

doc.add_heading('10.1. Monitoring Components', level=2)
monitoring = [
    ("Sentry", "Error Tracking", "Real-time error alerts, stack traces, user impact"),
    ("Datadog", "Performance Monitoring", "API response times, DB query perf, resource usage"),
    ("UptimeRobot", "Uptime Monitoring", "24/7 site availability checks"),
    ("Prometheus", "Metrics Collection", "System metrics (CPU, memory, DB)"),
    ("Grafana", "Visualization", "Monitoring dashboards"),
    ("Log Aggregation", "Log Management", "Structured logs from all services"),
    ("Alerting", "Notifications", "Slack/Email/PagerDuty alerts"),
]
add_table(["Tool", "Type", "Monitors"], monitoring)

doc.add_heading('10.2. Logging Structure', level=2)
add_code("""
LOG CATEGORIES AND WHAT IS LOGGED:

  API LOGS:      method, path, status, latency, user_id, IP
  AUTH LOGS:     login attempts, failures, OTP sent/verified
  ADMIN LOGS:    who changed what (audit), with old/new values
  AI LOGS:       prompts, responses, token usage, failures
  DB LOGS:       slow queries, connection errors, deadlocks
  SECURITY LOGS: rate limit hits, unauthorized access attempts
  CELERY LOGS:   task completion, task failures, retries

LOG FORMAT (JSON):
{
  "timestamp": "2026-09-03T10:30:00Z",
  "level": "ERROR",
  "service": "backend",
  "endpoint": "/api/v1/colleges",
  "method": "GET",
  "status": 500,
  "latency_ms": 234,
  "user_id": "u123",
  "ip": "1.2.3.4",
  "message": "Database connection timeout",
  "error": "..."
}
""")

doc.add_heading('10.3. Alert Thresholds', level=2)
add_table(["Alert", "Threshold", "Severity"], [
    ("API 5xx errors", "> 1% of requests in 5 min", "Critical"),
    ("API response time", "> 1000ms average in 5 min", "Warning"),
    ("Database CPU", "> 80% for 10 minutes", "Critical"),
    ("Database connections", "> 90% of max", "Warning"),
    ("API error rate", "Sudden spike", "Critical"),
    ("Disk space", "< 20% free", "Warning"),
    ("Memory usage", "> 85% for 10 min", "Critical"),
    ("Uptime", "Any downtime > 5 min", "Critical"),
    ("AI API failure", "> 5 failures in 5 min", "Warning"),
    ("Celery queue", "Task backlog > 100", "Warning"),
])

doc.add_page_break()

# ============================================
# 11. CACHING ARCHITECTURE
# ============================================
doc.add_heading('11. CACHING ARCHITECTURE (REDIS)', level=1)

doc.add_heading('11.1. What We Cache', level=2)
add_code("""
   CACHE STRATEGY
   ┌──────────────────────────────────────────────────────────┐
   │  L1: Browser Cache (Next.js ISR/SSG)                    │
   │  L2: CDN Cache (Cloudflare)                             │
   │  L3: Redis Cache (API responses)                        │
   │  L4: PostgreSQL (indexes + query optimization)          │
   └──────────────────────────────────────────────────────────┘

   REDIS KEY STRUCTURE:
   ┌──────────────────────────────────────────────┐
   │ cache:college:{id}          → college JSON   │
   │ cache:college:{id}:courses  → courses list   │
   │ cache:course:{id}           → course JSON    │
   │ cache:search:{hash}         → search results │
   │ cache:homepage              → homepage data  │
   │ cache:popular-courses       → popular list   │
   │ cache:featured-colleges     → featured list  │
   │ cache:sitemap               → sitemap XML    │
   │                            TTL: default 300s│
   └──────────────────────────────────────────────┘

   SESSION CACHE:
   session:{token}  → user session data

   RATE LIMITING:
   ratelimit:{ip}:{endpoint}  → request count + window
   ratelimit:{user}:{endpoint}→ request count + window

   CELERY BROKER:
   celery queue → task messages
""")

doc.add_heading('11.2. Cache Invalidation', level=2)
add_bullet("When college data is updated → clear cache:college:{id}")
add_bullet("When course data is updated → clear cache:course:{id}")
add_bullet("When any content updates → clear cache:homepage")
add_bullet("Search caches have short TTL (60s)")
add_bullet("Pattern: delete keys matching prefix after writes")

doc.add_page_break()

# ============================================
# 12. SEARCH ENGINE OPTIMIZATION (SEO) ARCHITECTURE
# ============================================
doc.add_heading('12. SEO ARCHITECTURE', level=1)

doc.add_heading('12.1. SEO Components', level=2)
add_bullet("Clean URLs: /colleges/bhms-colleges-in-karnataka")
add_bullet("Meta titles + descriptions for every page")
add_bullet("Canonical URLs to prevent duplicate content")
add_bullet("Open Graph + Twitter Card metadata")
add_bullet("Schema.org structured data (educational org, course, FAQ)")
add_bullet("XML Sitemaps (static + dynamic)")
add_bullet("robots.txt configuration")
add_bullet("Next.js SSR/SSG for server-rendered HTML")
add_bullet("Programmatic SEO pages (BHMS in Karnataka, etc.)")

doc.add_heading('12.2. Programmatic SEO Architecture', level=2)
add_code("""
   AUTOMATIC PAGE GENERATION
   ┌──────────────────────────────────────────────┐
   │ DATA:  Courses (BHMS, BAMS, MBBS...)         │
   │        × States (Karnataka, Bihar...)        │
   │        × Cities (Bangalore, Patna...)        │
   └──────────────┬───────────────────────────────┘
                  │
                  ▼
   ┌──────────────────────────────────────────────┐
   │ PAGE GENERATOR (Content API)                 │
   │   /colleges/bhms-colleges-in-karnataka       │
   │   /colleges/bams-colleges-in-karnataka       │
   │   /colleges/nursing-colleges-in-bihar        │
   │   /colleges/bpharm-colleges-in-bangalore     │
   └──────────────┬───────────────────────────────┘
                  │
                  ▼
   ┌──────────────────────────────────────────────┐
   │ QUALITY CHECK (Rule-based)                   │
   │  • Has >= 3 real colleges? → indexable       │
   │  • Has < 3 colleges?        → noindex        │
   │  • Duplicate content?       → canonical      │
   └──────────────────────────────────────────────┘
""")

doc.add_heading('12.3. Schema.org Markup', level=2)
add_code("""
  EDUCATIONAL ORGANIZATION (College page):
  {
    "@type": "EducationalOrganization",
    "name": "XYZ Medical College",
    "url": "https://padhaanewala.in/college/xyz",
    "address": { state, city, postalCode },
    "geo": { latitude, longitude },
    "aggregateRating": { ratingValue, reviewCount }
  }

  COURSE:
  {
    "@type": "Course",
    "name": "BHMS",
    "description": "...",
    "provider": { "@type": "Organization", "name": "College" },
    "hasCourseInstance": { duration }
  }

  FAQ:
  {
    "@type": "FAQPage",
    "mainEntity": [
      { "@type": "Question", "name": "...", "acceptedAnswer": {...} }
    ]
  }
""")

doc.add_page_break()

# ============================================
# 13. OFFICIAL CLOUD / HOSTING ARCHITECTURE
# ============================================
doc.add_heading('13. PRODUCTION HOSTING ARCHITECTURE', level=1)

add_code("""
                     INTERNET USERS
                          │
                          ▼
                 ┌─────────────────┐
                 │   CLOUDFLARE    │
                 │  CDN + WAF      │
                 │  SSL            │
                 └────────┬────────┘
                          │
                          ▼
                 ┌─────────────────┐
                 │   NGINX (LB)    │
                 │  Reverse Proxy  │
                 └────────┬────────┘
                          │
            ┌─────────────┼─────────────┐
            │             │             │
            ▼             ▼             ▼
     ┌────────────┐ ┌────────────┐ ┌────────────┐
     │  NEXT.JS   │ │  NEXT.JS   │ │  NEXT.JS   │
     │  instance 1│ │  instance 2│ │  instance 3│
     └────────────┘ └────────────┘ └────────────┘
            │             │             │
            └─────────────┼─────────────┘
                          │
                 ┌────────┴────────┐
                 │  NGINX (LB)    │
                 │  API Gateway   │
                 └────────┬────────┘
                          │
            ┌─────────────┼─────────────┐
            │             │             │
            ▼             ▼             ▼
     ┌────────────┐ ┌────────────┐ ┌────────────┐
     │  FASTAPI   │ │  FASTAPI   │ │  FASTAPI   │
     │  instance 1│ │  instance 2│ │  instance 3│
     │  (API)     │ │  (API)     │ │  (AI)      │
     └────────────┘ └────────────┘ └────────────┘
            │             │             │
            └──────┬──────┴──────┬──────┘
                   │             │
      ┌────────────┼─────────────┼──────────────┐
      ▼            ▼             ▼              ▼
┌──────────┐ ┌──────────┐ ┌──────────┐  ┌──────────────┐
│POSTGRESQL│ │  REDIS   │ │  CELERY  │  │ AWS S3 / R2  │
│  Primary │ │  Cache   │ │  Worker  │  │ Object Store │
└────┬─────┘ └──────────┘ └──────────┘  └──────────────┘
     │
     ▼
┌──────────┐
│POSTGRESQL│
│ Replica  │
└──────────┘

SERVICES (single host for small scale):
  FRONTEND:  3 Next.js instances
  BACKEND:   3 FastAPI instances
  DB:        PostgreSQL (1 primary, 1 replica)
  CACHE:     Redis (1 instance)
  WORKERS:   Celery (2-4 workers)
  PROXY:     Nginx (2 instances)

INFRA (recommended):
  - 2 application servers (active-active)
  - 1 database server
  - Load balancer in front
  - CDN on top (Cloudflare)
""")

doc.add_page_break()

# ============================================
# 14. USER JOURNEY ARCHITECTURE (Flows)
# ============================================
doc.add_heading('14. KEY USER JOURNEYS', level=1)

doc.add_heading('14.1. Student Discovery Journey', level=2)
add_code("""
   NEW VISITOR (Organic search / Google)
        │
        ▼
   SEARCHES "BHMS colleges in Karnataka"
        │
        ▼
   Lands on: /colleges/bhms-colleges-in-karnataka
        │
        ├── ▶ BROWSE college list (filter by fees, hostel)
        │         │
        │         ▼
        │    Opens a college page /college/xyz
        │         │
        │         ├── ▶ Reads overview, courses, fees
        │         ├── ▶ Sees reviews and gallery
        │         ├── ▶ Saves college (if logged in)
        │         └── ▶ Clicks "Get Admission Help"
        │
        ├── ▶ Uses AI Predictor /predictor
        │         │  (enters rank, course, budget)
        │         ▼
        │    Gets recommendations (Highly Suitable/Possible/Reach)
        │         │
        │         └── ▶ Clicks "Get Help" on a college
        │
        └── ▶ Uses AI Assistant /ai
                  │  (asks "What is BHMS?")
                  ▼
             Gets answer with sources
                  │
                  └── ▶ Converts to enquiry OR saves info
""")

doc.add_heading('14.2. Enquiry-to-Admission Journey', level=2)
add_code("""
   STUDENT                         ADMIN             COUNSELLOR
      │                                │                  │
      ▼                                │                  │
  Submits enquiry ────▶  API/POST /enquiries
      │                    │            │
      │                    ▼            │
      │                LEAD CREATED     │
      │                (New)            │
      │                    │            │
      │                    ▼            │
      │                NOTIFY admin     │
      │                + SMS/email ack  │
      │                    │            │
      │                    ▼            │
      │                ASSIGN counsellor│────────▶ Sees lead
      │                                 │            │
      │                                 │            ▼
      │                                 │        CONTACTS student
      │◀────────────────────────────────│───────── Calls/WhatsApp
      │                                 │            │
      │        Interested?              │            ▼
      │◀────────────────────────────────│───────── Updates: Interested
      │                                 │            │
      │                                 │            ▼
      │        Shares documents,        │     Helps with application
      │        guides through process   │            │
      │                                 │            ▼
      │                                 │     Updates: Application
      │◀────────────────────────────────│──────── Started
      │                                 │            │
      │                                 │            ▼
      │                                 │     Updates: Admission
      │                                 │     Completed / Not Int.
      │                                 │            │
      │                                 │            ▼
      │                                 │     LEAD CLOSED
      │                                 │
      └──(Lead source + UTM tracked
          from start for analytics)─────┘
""")

doc.add_heading('14.3. Mock Test Journey', level=2)
add_code("""
   LOGGED-IN STUDENT
        │
        ▼
   Visits /mock-tests
        │
        ▼
   Selects exam (e.g. NEET) + subject + difficulty
        │
        ▼
   Starts test (timer begins)
        │
        ├── ▶ Answers questions
        ├── ▶ Marks some for review
        ├── ▶ Navigates next/prev
        │
        ▼
   SUBMITS test
        │
        ▼
   RESULT PAGE:
     • Score: 85/100
     • Percentage: 85%
     • Correct: 75, Incorrect: 15, Unattempted: 10
     • Time taken: 45:32
     • Topic-wise performance: [Biology 90%, Physics 80%...]
     • Rank/percentile
        │
        ├── ▶ Practice Again
        ├── ▶ View Solutions
        └── ▶ View in dashboard history
""")

doc.add_page_break()

# ============================================
# 15. DATA FLOW ARCHITECTURE
# ============================================
doc.add_heading('15. DATA FLOW ARCHITECTURE', level=1)

doc.add_heading('15.1. Data Ingestion Paths', level=2)
add_code("""
   PATH 1: MANUAL ENTRY
   Admin → Admin Panel → Form → FastAPI → PostgreSQL

   PATH 2: CSV/EXCEL IMPORT
   Admin uploads CSV → FastAPI validates → Celery processes
   → Preview → Confirm → Import to PostgreSQL

   PATH 3: AUTOMATIC (Embeddings)
   Data change → Celery task → Chunk text → Embedding API
   → Store in document_embeddings (pgvector)

   PATH 4: IMAGE UPLOAD
   Admin uploads image → Validate → Optimize (Celery)
   → Upload to S3/R2 → Store URL in PostgreSQL → CDN serves

   PATH 5: ENQUIRY
   Student submits form → FastAPI validates → Save to DB
   → Notify admin (email/SMS) → Assign counsellor
""")

doc.add_heading('15.2. Data Read Paths', level=2)
add_code("""
   PUBLIC WEBSITE (SEO pages):
   Browser → CDN → Next.js (SSG/ISR) → PostgreSQL (cached)

   COLLEGE PAGE:
   Browser → CDN → Next.js → Redis cache (hit?)
   → miss → FastAPI → PostgreSQL → cache → response

   SEARCH:
   Browser → Next.js → FastAPI → Rate limit → Cache (query hash)
   → miss → PostgreSQL (filter + vector search) → cache → response

   AI ASSISTANT:
   Browser → Next.js → FastAPI → Validate → RAG Pipeline
   → DB + vector search → LLM API → response → log usage

   ADMIN PANEL:
   Browser (auth) → Next.js → FastAPI (auth+RBAC) → PostgreSQL
""")

doc.add_page_break()

# ============================================
# 16. ROLES & PERMISSIONS MATRIX
# ============================================
doc.add_heading('16. ROLE-BASED ACCESS CONTROL (RBAC) MATRIX', level=1)

permissions = [
    ("Homepage/CMS", "Read/Write", "Read", "No", "Read/Write", "Read"),
    ("Colleges", "Read/Write", "Read/Write", "No", "No", "Read"),
    ("Courses", "Read/Write", "Read/Write", "No", "No", "Read"),
    ("Scholarships", "Read/Write", "Read/Write", "No", "No", "Read"),
    ("Exams", "Read/Write", "Read/Write", "No", "No", "Read"),
    ("Mock Tests", "Read/Write", "No", "No", "Read/Write", "No"),
    ("Questions", "Read/Write", "No", "No", "Read/Write", "No"),
    ("Students", "Read/Write", "No", "Assigned only", "No", "No"),
    ("Reviews", "Read/Write", "Moderate", "No", "No", "Read"),
    ("Blogs", "Read/Write", "Read/Write", "No", "No", "No"),
    ("Leads", "Read/Write", "Read", "Assigned only", "No", "No"),
    ("Counsellors", "Read/Write", "No", "No", "No", "No"),
    ("Media", "Read/Write", "Read/Write", "No", "No", "No"),
    ("SEO", "Read/Write", "No", "No", "No", "Read/Write"),
    ("Settings", "Read/Write", "No", "No", "No", "No"),
    ("Audit Logs", "Read/Write", "Read", "No", "No", "Read"),
    ("Notifications", "Read/Write", "Read/Write", "Read", "Read/Write", "Read"),
    ("Users", "Read/Write", "Read", "No", "No", "Read"),
]
add_table(["Module", "Super Admin", "Content Admin", "Counsellor", "Test Admin", "SEO Admin"], permissions)

doc.add_page_break()

# ============================================
# 17. API ARCHITECTURE LAYERS
# ============================================
doc.add_heading('17. API ARCHITECTURE', level=1)

add_code("""
   REQUEST LIFECYCLE
   ┌─────────────────────────────────────────────────────────────┐
   │  CLIENT (Next.js)                                           │
   └──────────────────────────────┬──────────────────────────────┘
                                  │ HTTPS
                                  ▼
   ┌─────────────────────────────────────────────────────────────┐
   │  MIDDLEWARES                                                │
   │  1. CORS                                                  │
   │  2. Rate Limit (Redis)                                    │
   │  3. Security Headers                                      │
   │  4. Logging                                               │
   │  5. Auth (JWT verify if protected)                        │
   └──────────────────────────────┬──────────────────────────────┘
                                  │
                                  ▼
   ┌─────────────────────────────────────────────────────────────┐
   │  ROUTER LAYER                                              │
   │  /api/v1/{module}/{action}/{id}                            │
   └──────────────────────────────┬──────────────────────────────┘
                                  │
                                  ▼
   ┌─────────────────────────────────────────────────────────────┐
   │  SCHEMA VALIDATION (Pydantic)                              │
   │  Request body → validate → sanitize                        │
   └──────────────────────────────┬──────────────────────────────┘
                                  │
                                  ▼
   ┌─────────────────────────────────────────────────────────────┐
   │  SERVICE LAYER (Business logic)                            │
   │  + Authorization check (RBAC)                              │
   └──────────────────────────────┬──────────────────────────────┘
                                  │
                                  ▼
   ┌─────────────────────────────────────────────────────────────┐
   │  REPOSITORY LAYER (SQLAlchemy)                             │
   │  + SQL injection prevention (ORM)                          │
   └──────────────────────────────┬──────────────────────────────┘
                                  │
                                  ▼
   ┌─────────────────────────────────────────────────────────────┐
   │  POSTGRESQL / REDIS                                        │
   └─────────────────────────────────────────────────────────────┘
              │
              └──▶ Response → JSON → Client
""")

doc.add_heading('17.1. API Versioning & Documentation', level=2)
add_bullet("All endpoints under /api/v1/")
add_bullet("FastAPI automatic OpenAPI docs at /docs")
add_bullet("Swagger UI for interactive testing")
add_bullet("Standard response format: { success, data, message, error }")
add_bullet("Standard error codes: 400 bad request, 401 unauthenticated, 403 forbidden, 404 not found, 429 rate limit, 500 server error")

doc.add_heading('17.2. Standard API Response Format', level=2)
add_code("""
SUCCESS:
{
  "success": true,
  "data": { ... },
  "message": "College fetched successfully",
  "error": null
}

ERROR:
{
  "success": false,
  "data": null,
  "message": "User-friendly error message",
  "error": {
    "code": "VALIDATION_ERROR",
    "details": "Field 'email' must be valid"
  }
}

PAGINATION:
{
  "success": true,
  "data": [ ... ],
  "meta": {
    "page": 1,
    "limit": 20,
    "total": 500,
    "total_pages": 25,
    "has_next": true,
    "has_prev": false
  }
}
""")

# ============================================
# CLOSING
# ============================================
doc.add_page_break()
add_title("PADHAANEWALA EDUTECH SERVICES", size=14)
add_title("BENGALURU - 560100", size=12, bold=False)
add_title("End of Complete System Architecture Document", size=12, bold=False)

output_path = r"D:\code\Clients\Padhaanewala\Padhaanewala_Complete_Architecture.docx"
doc.save(output_path)
print(f"Architecture document saved to: {output_path}")
